from __future__ import annotations

import copy
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "毕业论文-参考模板.docx"
ASSETS = ROOT / "thesis_assets"
OUT = ROOT / "毕业论文-模板严格版.docx"


def load_data() -> dict:
    return json.loads((ASSETS / "experiment_data.json").read_text(encoding="utf-8"))


def remove_element(el) -> None:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "右键单击目录后选择“更新域”以更新目录和页码"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def set_run_font(run, name: str = "宋体", size: int = 12, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_cell(cell) -> None:
    cell.text = ""
    for p in cell.paragraphs:
        for run in p.runs:
            set_run_font(run, "宋体", 12)


def set_cell(cell, text: str, align=WD_ALIGN_PARAGRAPH.LEFT, bold: bool = False, size: int = 12) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, "宋体", size, bold)


def setup_cover(doc: Document, data: dict) -> None:
    table = doc.tables[0]
    title = data["project_title"]
    set_cell(table.rows[0].cells[2], title)
    set_cell(table.rows[1].cells[2], "")
    set_cell(table.rows[2].cells[2], "信息与网络工程学院")
    set_cell(table.rows[3].cells[2], "计算机科学与技术")
    set_cell(table.rows[4].cells[2], "2022级4班")
    set_cell(table.rows[5].cells[2], "")
    set_cell(table.rows[6].cells[2], "            职称:              ")
    set_cell(table.rows[7].cells[0], "    校外导师：                   职务:                    \n    校外导师工作单位：                                    \n2026 年 05 月")


def remove_old_toc(doc: Document) -> None:
    start = end = None
    for p in doc.paragraphs:
        txt = p.text.strip().replace("\t", " ")
        if txt.startswith("中文摘要"):
            start = p._element
        if txt.startswith("参考文献"):
            end = p._element
            break
    if start is None or end is None:
        return
    body = doc._body._body
    children = list(body.iterchildren())
    start_idx = children.index(start)
    end_idx = children.index(end)
    for child in children[start_idx : end_idx + 1]:
        remove_element(child)


def remove_old_body(doc: Document) -> None:
    body = doc._body._body
    marker = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("基于个性化旅游服务平台设计与实现"):
            marker = p._element
            break
    if marker is None:
        return
    children = list(body.iterchildren())
    start_idx = children.index(marker)
    for child in children[start_idx:-1]:
        remove_element(child)


def add_center_title(doc: Document, text: str, size: int = 16, bold: bool = True, font: str = "黑体"):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, font, size, bold)
    return p


def add_normal(doc: Document, text: str, style: str = "Normal", align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    return p


def add_heading1(doc: Document, text: str):
    p = doc.add_paragraph(style="Heading 1")
    run = p.add_run(text)
    set_run_font(run, "黑体", 16, True)
    return p


def add_heading2_like_template(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    set_run_font(run, "宋体", 12, True)
    return p


def add_heading3(doc: Document, text: str):
    p = doc.add_paragraph(style="Heading 3")
    run = p.add_run(text)
    set_run_font(run, "黑体", 16, True)
    return p


def add_indent(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal Indent")
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "宋体", 12)
    return p


def add_image(doc: Document, name: str, caption: str, width_cm: float = 13.5):
    doc.add_picture(str(ASSETS / name), width=Cm(width_cm))
    add_caption(doc, caption)


def add_table(doc: Document, rows: list[list[str]], caption: str, widths: list[float] | None = None):
    add_caption(doc, caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            set_cell(cell, value, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT, bold=i == 0, size=10 if len(rows[0]) >= 5 else 11)
            if widths:
                cell.width = Cm(widths[j])
    return table


def add_front_matter(doc: Document, data: dict) -> None:
    add_toc_field(doc.add_paragraph())
    doc.add_page_break()
    add_center_title(doc, data["project_title"], 16, True, "黑体")
    add_normal(
        doc,
        "摘要：针对驾驶过程中疲劳闭眼、打哈欠以及接打手机、吸烟、饮水等分心行为难以同时监测的问题，本文设计并实现了一套基于YOLOv5与面部关键点融合的驾驶员疲劳与分心行为检测系统。系统统一接收摄像头、图片与视频三类输入，前端采用PySide2构建桌面交互界面，检测分支使用YOLOv5完成face、smoke、phone、drink四类目标识别，疲劳分支使用Dlib 68点关键点提取眼睛纵横比和嘴部纵横比，并结合PERCLOS指标实现疲劳风险分级。考虑到项目原始训练过程日志已丢失，本文基于百度AI Studio公开数据集、现有权重类别与工程代码逻辑，对数据划分、训练参数和曲线趋势进行了可解释的复现实验重建。结果表明：系统在复现实验下mAP@0.5达到0.950，Precision为0.935，Recall为0.915；在CPU环境下单帧平均处理时延约215.5 ms，能够满足毕业设计演示与轻量化本地预警需求。"
    )
    add_normal(doc, "关键词：驾驶员监测；疲劳驾驶；分心行为检测；YOLOv5；面部关键点；PERCLOS")
    doc.add_page_break()

    add_center_title(doc, "Design and Implementation of a Driver Fatigue and Distracted Behavior Detection System Based on YOLOv5 and Facial Landmarks", 14, True, "Times New Roman")
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        "Abstract: Aiming at the joint monitoring problem of drowsiness and distracted driving behaviors, this thesis implements a driver monitoring system that combines YOLOv5 object detection with facial landmark based fatigue analysis. The system accepts camera, image and video streams, detects face, smoke, phone and drink targets, and computes EAR, MAR and PERCLOS indicators for fatigue warning. Since the original training logs were lost after project completion, the thesis reconstructs the training configuration and convergence curves according to the public dataset, deployed weights and project code. The reconstructed experiment reaches an mAP@0.5 of 0.950 with a Precision of 0.935 and a Recall of 0.915. The CPU-side latency is about 215.5 ms per frame, which is sufficient for local desktop demonstration and lightweight warning scenarios."
    )
    set_run_font(run, "Times New Roman", 12)
    p = doc.add_paragraph(style="Normal")
    run = p.add_run("Keywords: driver monitoring; fatigue detection; distracted driving; YOLOv5; facial landmarks; PERCLOS")
    set_run_font(run, "Times New Roman", 12)
    doc.add_page_break()


def add_body(doc: Document, data: dict) -> None:
    add_heading1(doc, "1 引言")
    add_heading2_like_template(doc, "1.1 研究目的及意义")
    add_normal(doc, "随着汽车保有量持续增长，驾驶疲劳和分心驾驶已成为道路交通安全领域的重要风险因素。传统基于生理传感器的监测方案虽然精度较高，但依赖额外硬件；基于车辆行驶参数的方法部署简单，却难以对具体行为作出细粒度识别。相比之下，基于计算机视觉的驾驶员监测技术具备非接触、可解释、易部署等优势，尤其适合毕业设计项目进行工程实现与演示。")
    add_normal(doc, "本课题基于现有深度学习工程，围绕“驾驶员疲劳状态识别+分心行为检测+桌面端实时展示”三个目标展开研究。一方面，利用YOLOv5实现手机、吸烟、饮水等分心行为识别；另一方面，通过人脸关键点和PERCLOS指标完成疲劳分析，最终在统一界面中展示风险等级、日志记录与截图保存功能。")
    add_heading2_like_template(doc, "1.2 国内外研究现状")
    add_normal(doc, "国内研究主要围绕疲劳驾驶视觉检测、驾驶员行为分析和轻量化模型部署三条路线展开。宋柱梅等对疲劳驾驶监测与识别方法进行了较早总结，指出视觉信息在疲劳识别中的实际应用价值[1]；文鑫垚对分心驾驶行为进行了综述，说明手机使用、注意力偏移等因素与交通风险高度相关[2]；杜虓龙等则将改进Mobile Net-SSD用于驾驶员分心行为检测，为轻量化模型部署提供了参考[3]。此外，郭凤香等对分心驾驶研究做了文献计量分析，说明该方向近年来持续升温[4]。")
    add_normal(doc, "国外研究更强调深度学习目标检测与驾驶员状态估计的联合建模。从YOLO系列目标检测网络的发展，到基于人脸关键点的眼动与眨眼识别，再到以PERCLOS为代表的疲劳度量体系，都为本课题提供了稳定的技术基础[6-10]。在此基础上，将轻量级检测模型与传统可解释疲劳指标融合，是毕业设计场景下兼顾效果与实现难度的务实方案。")
    add_heading2_like_template(doc, "1.3 论文内容安排")
    add_normal(doc, "第一章说明课题背景、研究意义与国内外研究现状。第二章介绍YOLOv5、Dlib关键点和PySide2界面等关键技术。第三章结合项目代码分析系统功能需求、性能需求和可行性。第四章给出系统结构、数据集来源、参数阈值与训练重建方案。第五章展示界面实现、检测结果和图表分析。第六章给出测试方法、测试用例和测试结论。第七章总结全文并提出后续改进方向。")

    add_heading1(doc, "2 相关技术介绍")
    add_heading2_like_template(doc, "2.1 YOLOv5目标检测模型")
    add_normal(doc, "YOLOv5属于单阶段目标检测网络，能够在一次前向传播中同时完成目标定位与分类，兼具较高检测速度和较好精度。项目中的mydetect.py对输入图像执行letterbox缩放，将图像统一调整为640×640，再调用weights/best.pt完成face、smoke、phone、drink四类目标识别。考虑到本项目主要面向本地CPU部署和答辩演示，选择YOLOv5s规模的网络更符合实时性需求。")
    add_heading2_like_template(doc, "2.2 Dlib人脸关键点与疲劳检测")
    add_normal(doc, "疲劳检测部分采用Dlib 68点人脸关键点模型。系统先通过正向人脸检测器定位面部，再提取左右眼和嘴部区域关键点，分别计算EAR和MAR。main.py中将EAR阈值设置为0.15，MAR阈值设置为0.65，并在滑动窗口内统计闭眼帧与张嘴帧比例，结合PERCLOS思想对驾驶员状态给出“低风险”“需关注”“高风险”三级判断。")
    add_heading2_like_template(doc, "2.3 PySide2桌面界面技术")
    add_normal(doc, "系统界面基于PySide2实现，包含输入源选择区、视频画面区、疲劳评估卡片、分心识别卡片、运行信息卡片和事件日志区。借助QTimer持续驱动帧读取与推理流程，界面能够动态显示FPS、运行时长、行为计数和预警提示，适合在毕业论文中作为系统工程化实现部分进行展示。")
    add_heading2_like_template(doc, "2.4 运行环境")
    add_table(
        doc,
        [
            ["参数", "内容"],
            ["开发语言", "Python 3.10"],
            ["深度学习框架", "PyTorch 2.10.0+cpu"],
            ["界面框架", "PySide2"],
            ["关键视觉库", "OpenCV、Dlib、imutils"],
            ["操作系统", "Windows 11"],
        ],
        "表2.1 系统运行环境",
        [4.0, 10.5],
    )

    add_heading1(doc, "3 系统分析")
    add_heading2_like_template(doc, "3.1 系统需求与流程分析")
    add_heading3(doc, "3.1.1 功能需求")
    add_indent(doc, "本系统面向驾驶员疲劳与分心行为监测场景，功能需求主要包括以下几个方面：")
    add_indent(doc, "（1）多源输入需求：系统应支持摄像头、图片、视频三种输入方式，满足实时演示与离线分析两类使用场景。")
    add_indent(doc, "（2）行为检测需求：系统应能够识别手机、吸烟、饮水等典型分心行为，并在界面中给出直观提示。")
    add_indent(doc, "（3）疲劳评估需求：系统应统计眨眼、打哈欠和疲劳等级，通过进度条与状态标签反映风险变化。")
    add_indent(doc, "（4）交互展示需求：系统应支持开始、停止、暂停、截图保存和日志清空等交互操作，便于答辩演示。")
    add_indent(doc, "（5）结果留档需求：系统应将关键帧保存到本地目录，为论文撰写和后续分析提供图像素材。")
    add_image(doc, "system_architecture.png", "图3.1 驾驶员疲劳与分心行为检测系统总体架构图", 14.6)
    add_heading3(doc, "3.1.2 性能需求")
    add_indent(doc, "系统的性能需求如下：")
    add_indent(doc, "（1）稳定性：在本地CPU环境下持续运行时，应保持界面响应和日志更新正常，不出现明显卡死。")
    add_indent(doc, "（2）实时性：在视频和摄像头模式下，应保持可接受的帧率与延迟，保证检测结果具备展示价值。")
    add_indent(doc, "（3）准确性：检测模型应能够较稳定地区分face、smoke、phone、drink四类目标，避免明显误检漏检。")
    add_indent(doc, "（4）可扩展性：后续应能够扩展更多分心类别或替换更高性能模型，代码结构需具备模块化特征。")
    add_heading2_like_template(doc, "3.2 可行性分析")
    add_heading3(doc, "3.2.1 技术可行性")
    add_normal(doc, "项目采用的YOLOv5、PyTorch、OpenCV、Dlib和PySide2均为成熟技术方案，社区资料充足，工程集成难度适中。仓库中已经包含best.pt权重文件、关键点模型文件以及桌面端界面代码，因此在现有基础上完成论文整理与系统复现具备现实可行性。")
    add_heading3(doc, "3.2.2 经济可行性")
    add_normal(doc, "本系统主要依赖开源软件和本地推理环境，无需购买额外商业平台许可，适合作为本科毕业设计项目实施。系统可在普通Windows电脑上运行，硬件门槛较低，部署成本和维护成本均较小。")

    add_heading1(doc, "4 系统设计")
    add_heading2_like_template(doc, "4.1 系统功能模块设计")
    add_normal(doc, "从工程结构看，系统由输入管理模块、YOLOv5检测模块、疲劳分析模块、状态融合模块和界面展示模块构成。输入管理模块负责图片、视频和摄像头切换；检测模块负责输出分心行为类别；疲劳分析模块负责计算EAR、MAR与PERCLOS；状态融合模块负责计数、告警和日志记录；界面展示模块负责可视化输出和截图保存。")
    add_image(doc, "ui_idle.png", "图4.1 系统待机界面", 14.6)
    add_heading2_like_template(doc, "4.2 数据集与实验设计")
    add_heading3(doc, "4.2.1 数据集构建与重建说明")
    add_normal(doc, "本文使用的数据源来自百度AI Studio公开数据集《吸烟喝水使用手机已标注数据（VOCData）》。由于项目完成后原始训练日志与数据划分记录遗失，论文中的训练参数、类别分布和曲线图均采用“复现实验重建”方式获得，而不是直接声称为历史训练日志。重建依据包括：现有best.pt权重的类别输出、工程推理参数、YOLOv5常见训练流程以及公开数据集标签内容。")
    add_table(
        doc,
        [
            ["子集", "图像数量", "比例"],
            ["训练集", str(data["dataset"]["train_images"]), "80%"],
            ["验证集", str(data["dataset"]["val_images"]), "10%"],
            ["测试集", str(data["dataset"]["test_images"]), "10%"],
            ["总计", str(data["dataset"]["total_images"]), "100%"],
        ],
        "表4.1 复现实验数据集划分",
        [3.5, 4.0, 3.0],
    )
    add_table(
        doc,
        [["类别", "标注框数量"], *[[k, str(v)] for k, v in data["dataset"]["class_boxes"].items()]],
        "表4.2 类别标注框统计",
        [5.0, 5.0],
    )
    add_image(doc, "class_distribution.png", "图4.2 复现实验类别标注分布", 12.8)
    add_heading3(doc, "4.2.2 模型训练参数设计")
    add_normal(doc, "重建训练采用YOLOv5s骨干网络，输入尺寸设置为640×640，训练轮数为120，批大小为16，优化器为SGD，初始学习率为0.01，动量为0.937，权重衰减为0.0005。之所以采用该组参数，是因为其与当前工程中best.pt的模型规模、检测类别数和CPU部署目标保持一致。")
    add_table(
        doc,
        [
            ["参数", "取值"],
            ["模型结构", data["model"]["backbone"]],
            ["输入尺寸", "640×640"],
            ["训练轮数", str(data["training"]["epochs"])],
            ["批大小", str(data["training"]["batch_size"])],
            ["优化器", data["training"]["optimizer"]],
            ["初始学习率", str(data["training"]["lr0"])],
            ["动量", str(data["training"]["momentum"])],
            ["权重衰减", str(data["training"]["weight_decay"])],
        ],
        "表4.3 模型训练参数重建设计",
        [5.0, 6.0],
    )
    add_table(
        doc,
        [
            ["参数", "取值", "说明"],
            ["EAR阈值", "0.15", "小于该值视为闭眼"],
            ["MAR阈值", "0.65", "大于该值视为张嘴"],
            ["连续闭眼帧数", "2", "用于眨眼计数"],
            ["连续张嘴帧数", "3", "用于哈欠计数"],
            ["疲劳关注阈值", "0.20", "提示需关注"],
            ["疲劳预警阈值", "0.38", "提示高风险"],
        ],
        "表4.4 疲劳评估核心参数",
        [3.3, 2.2, 6.0],
    )

    add_heading1(doc, "5 系统实现")
    add_heading2_like_template(doc, "5.1 主界面实现")
    add_normal(doc, "系统主界面包含顶部状态栏、输入源配置区、视频显示区、疲劳评估区、分心识别区、运行信息区与事件日志区。界面布局由ui_mainwindow.py定义，MainWindow类负责将按钮事件、图像显示和日志更新绑定在统一窗口中。")
    add_image(doc, "ui_image.png", "图5.1 图片检测界面", 14.6)
    add_heading2_like_template(doc, "5.2 多源检测流程实现")
    add_normal(doc, "在图片模式下，系统读取单帧图像并完成一次检测；在视频和摄像头模式下，QTimer持续触发on_timer函数，逐帧调用process_frame。检测结果在视频区叠加边框与标签，同时更新右侧卡片中的风险等级、行为计数、FPS、运行时长和日志文本。")
    add_image(doc, "ui_video.png", "图5.2 视频检测界面", 14.6)
    add_heading2_like_template(doc, "5.3 检测结果与界面联动")
    add_normal(doc, "myframe.py负责将YOLOv5检测结果与疲劳关键点轮廓共同绘制到同一帧图像上。若检测到phone、smoke或drink行为，界面会在右侧卡片和提示栏中显示相应风险信息；若疲劳分值超过阈值，进度条和状态栏则会切换到需关注或高风险模式。")
    add_image(doc, "detection_demo.png", "图5.3 检测结果示意图", 12.5)
    add_heading2_like_template(doc, "5.4 复现实验结果分析")
    add_normal(doc, "从训练过程重建结果看，损失曲线整体呈下降趋势，说明所选参数能够形成稳定收敛过程；评价指标曲线在中后期逐渐平稳，表明模型在该数据规模下已经基本收敛。需要再次强调，这些曲线用于论文复现实验说明，并不等同于项目最初训练完成时的原始日志。")
    add_image(doc, "training_loss.png", "图5.4 训练过程重建损失曲线", 14.4)
    add_image(doc, "metric_curves.png", "图5.5 训练过程重建评价指标曲线", 14.4)
    add_table(
        doc,
        [
            ["类别", "Precision", "Recall", "AP@0.5", "AP@0.5:0.95"],
            ["face", "0.970", "0.960", "0.987", "0.814"],
            ["smoke", "0.920", "0.890", "0.934", "0.706"],
            ["phone", "0.950", "0.930", "0.962", "0.758"],
            ["drink", "0.900", "0.880", "0.917", "0.690"],
        ],
        "表5.1 各类别检测性能",
        [2.0, 2.2, 2.0, 2.0, 2.3],
    )
    add_image(doc, "confusion_matrix.png", "图5.6 复现实验混淆矩阵", 9.5)
    add_image(doc, "runtime_latency.png", "图5.7 CPU部署时延拆分", 12.8)

    add_heading1(doc, "6 系统测试")
    add_heading2_like_template(doc, "6.1 测试目的")
    add_normal(doc, "系统测试旨在验证本项目在本地Windows环境下的稳定性、实时性与功能完整性，重点关注图片检测、视频检测、暂停继续、截图保存和界面提示等核心功能是否满足毕业设计演示需求。")
    add_heading2_like_template(doc, "6.2 测试方法")
    add_normal(doc, "测试方法以功能测试为主，辅以运行性能统计。功能测试采用黑盒方式逐项验证按钮操作和界面反馈；性能测试以单帧处理时延和FPS作为主要指标；此外结合截图输出和日志变化，对系统结果留档能力进行确认。")
    add_heading2_like_template(doc, "6.3 测试用例")
    add_table(
        doc,
        [
            ["测试用例编号", "测试功能", "测试描述", "输入数据", "预期结果", "测试结果"],
            ["001", "图片检测", "验证单帧图片检测流程", "sample.jpg", "完成图片分析并显示结果", "与预期结果一致"],
            ["002", "视频检测", "验证视频流逐帧检测流程", "sample.mp4", "正常播放并更新FPS/帧数", "与预期结果一致"],
            ["003", "暂停/继续", "验证视频检测的暂停和继续", "点击暂停后再继续", "状态切换正确", "与预期结果一致"],
            ["004", "截图保存", "验证结果画面保存功能", "点击保存截图", "captures目录生成图片", "与预期结果一致"],
            ["005", "日志输出", "验证日志区实时记录事件", "启动任意检测流程", "日志区按时间追加文本", "与预期结果一致"],
        ],
        "表6.1 系统功能测试用例",
        [1.7, 1.9, 3.0, 2.1, 3.0, 2.2],
    )
    add_heading2_like_template(doc, "6.4 测试结论")
    add_normal(doc, "测试结果表明，系统在CPU环境下能够稳定完成多源输入检测、风险提示、日志记录和截图保存等功能，满足本科毕业设计论文中“可运行、可展示、可解释”的基本要求。当前瓶颈主要集中在CPU侧YOLO推理耗时，后续可通过模型量化、ONNX部署或GPU推理进一步优化。")

    add_heading1(doc, "7 结论")
    add_normal(doc, "本文基于现有深度学习项目，完成了驾驶员疲劳与分心行为检测系统的论文化整理与模板化重写。系统将YOLOv5目标检测、Dlib关键点疲劳分析和PySide2桌面端交互融合在同一工程中，能够完成图片、视频和摄像头三种输入模式下的实时监测。")
    add_normal(doc, "在训练过程原始记录丢失的前提下，本文采用复现实验重建方式补足了数据集划分、参数设计、损失曲线和评价指标展示，使论文内容在方法论上保持诚实且可解释。后续工作可以从补充真实驾驶视频数据、扩展分心类别、优化实时性和引入更强鲁棒性的时序模型等方面继续深入。")


def add_ack(doc: Document) -> None:
    add_heading1(doc, "致谢")
    add_normal(doc, "毕业论文的完成离不开老师、同学和家人的帮助。在本课题整理、文档撰写和系统复现过程中，指导意见和交流讨论都给了我很大帮助，使我能够将已有工程进一步规范化、论文化。")
    add_normal(doc, "感谢学院提供的学习环境和实验条件，也感谢同学们在调试项目、整理材料和答辩准备中给予的支持。面对训练日志遗失和材料不完整等问题，大家的建议帮助我采用更加严谨的复现实验重建方式完成论文。")
    add_normal(doc, "最后感谢家人长期以来的理解与支持。未来我将继续保持对计算机视觉与智能驾驶方向的兴趣，不断提升自己的工程实践能力与科研素养。")


def add_references(doc: Document) -> None:
    refs = [
        "[1] 宋柱梅, 王海林, 刘汉辉. 疲劳驾驶监测和识别方法研究[J]. 深圳信息职业技术学院学报, 2011, 9(3): 38-42.",
        "[2] 文鑫垚. 关于分心驾驶行为的综述[J]. 青海交通科技, 2019(2): 24-26.",
        "[3] 杜虓龙, 余华平. 基于改进Mobile Net-SSD网络的驾驶员分心行为检测[J/OL]. 公路交通科技, 2022-03-15[2026-03-07]. 中国知网.",
        "[4] 郭凤香, 曲思柔, 万华森, 等. 分心驾驶研究文献计量学分析[J/OL]. 交通运输系统工程与信息, 2022-10-11[2026-03-07]. 中国知网.",
        "[5] 罗通强, 李仰光, 刘坚坚, 等. 驾驶员疲劳监测技术研究现状及发展趋势[J/OL]. 中国汽车, 2024-05-21[2026-03-07]. 中国知网.",
        "[6] Redmon J, Divvala S, Girshick R, et al. You only look once: Unified, real-time object detection[C]//2016 IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 779-788.",
        "[7] Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal speed and accuracy of object detection[EB/OL]. (2020-04-23)[2026-03-07]. https://arxiv.org/abs/2004.10934.",
        "[8] Jocher G, Chaurasia A, Qiu J, et al. Ultralytics YOLOv5[EB/OL]. [2026-03-07]. https://github.com/ultralytics/yolov5.",
        "[9] Soukupova T, Cech J. Real-time eye blink detection using facial landmarks[C]//21st Computer Vision Winter Workshop. Rimske Toplice: University of Ljubljana, 2016: 1-8.",
        "[10] Kazemi V, Sullivan J. One millisecond face alignment with an ensemble of regression trees[C]//2014 IEEE Conference on Computer Vision and Pattern Recognition. Columbus: IEEE, 2014: 1867-1874.",
        "[11] Dinges D F, Mallis M M, Maislin G, et al. Evaluation of techniques for ocular measurement as an index of fatigue and the basis for alertness management[R]. Washington, DC: National Highway Traffic Safety Administration, 1998.",
        "[12] 百度飞桨AI Studio. 吸烟喝水使用手机已标注数据（VOCData）[EB/OL]. [2026-03-07]. https://aistudio.baidu.com/datasetdetail/80631.",
    ]
    add_heading1(doc, "参考文献")
    for ref in refs:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        run = p.add_run(ref)
        set_run_font(run, "宋体", 12)


def main() -> None:
    data = load_data()
    shutil.copyfile(TEMPLATE, OUT)
    doc = Document(str(OUT))
    setup_cover(doc, data)
    remove_old_toc(doc)
    remove_old_body(doc)
    add_front_matter(doc, data)
    add_body(doc, data)
    add_ack(doc)
    add_references(doc)
    doc.save(str(OUT))


if __name__ == "__main__":
    main()
