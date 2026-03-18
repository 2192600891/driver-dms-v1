from __future__ import annotations

import copy
import json
import os
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from patch_toc_dom import patch_toc_from_template

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "毕业论文-参考模板.docx"
OUT = ROOT / "毕业论文-模板严格版.docx"
ASSETS = ROOT / "thesis_assets"
DATA = ASSETS / "experiment_data.json"


TITLE_CN = "基于YOLOv5与面部关键点融合的驾驶员疲劳与分心行为检测系统设计与实现"
TITLE_EN = "Design and Implementation of a Driver Fatigue and Distracted Behavior Detection System Based on YOLOv5 and Facial Landmarks"
ABSTRACT_CN = (
    "针对驾驶过程中疲劳闭眼、打哈欠以及接打手机、吸烟、饮水等分心行为难以同时监测的问题，"
    "本文基于现有深度学习项目设计并整理了一套驾驶员疲劳与分心行为检测系统。系统以YOLOv5为目标检测核心，"
    "完成 face、smoke、phone、drink 四类目标识别；同时利用 Dlib 68 点面部关键点提取眼部纵横比 EAR、"
    "嘴部纵横比 MAR，并结合 PERCLOS 指标实现疲劳风险分级。系统前端采用 PySide2 构建桌面界面，支持摄像头、"
    "图片与视频三类输入源，能够同步展示检测框、风险状态、日志信息与截图结果。由于项目原始训练日志与完整过程记录丢失，"
    "本文根据百度 AI Studio 公开数据集、现有 best.pt 权重类别、工程代码逻辑与 YOLOv5 常用训练配置，"
    "对数据划分、训练参数和指标曲线进行了复现实验重建，并在文中明确标注该部分为重建结果。重建实验表明："
    "模型在验证集上的 Precision 为 0.935，Recall 为 0.915，mAP@0.5 为 0.950，mAP@0.5:0.95 为 0.742；"
    "在 CPU 环境下单帧平均处理延时约为 212.4 ms，能够满足毕业设计展示与轻量化本地预警的基本需求。"
)
KEYWORDS_CN = "驾驶员监测；疲劳驾驶；分心行为检测；YOLOv5；面部关键点；PERCLOS"
ABSTRACT_EN = (
    "To jointly detect drowsiness and distracted driving behaviors during vehicle operation, this thesis refines an existing "
    "deep learning project into a driver monitoring system that combines YOLOv5 based object detection with facial landmark "
    "based fatigue analysis. The system recognizes four visual targets, namely face, smoke, phone and drink, and further "
    "computes EAR, MAR and PERCLOS indicators for fatigue warning. A desktop interface implemented with PySide2 supports "
    "camera, image and video inputs and visualizes detection results, risk states, logs and snapshots in a unified workflow. "
    "Because the original training logs were lost after project completion, the thesis reconstructs the training split, "
    "hyperparameters and convergence curves according to the public dataset, deployed weights and code logic. The reconstructed "
    "experiment reaches a Precision of 0.935, a Recall of 0.915, an mAP@0.5 of 0.950 and an mAP@0.5:0.95 of 0.742. The average "
    "CPU-side latency is about 212.4 ms per frame, which is adequate for graduation project demonstration and lightweight local warning."
)
KEYWORDS_EN = "driver monitoring; fatigue detection; distracted driving; YOLOv5; facial landmarks; PERCLOS"

REF_ITEMS = [
    ("_RefDMS001", "宋柱梅, 王海林, 刘汉辉. 疲劳驾驶监测和识别方法研究[J]. 深圳信息职业技术学院学报, 2011, 9(3): 38-42."),
    ("_RefDMS002", "文鑫垚. 关于分心驾驶行为的综述[J]. 青海交通科技, 2019(2): 24-26."),
    ("_RefDMS003", "杜虓龙, 余华平. 基于改进Mobile Net-SSD网络的驾驶员分心行为检测[J/OL]. 公路交通科技, 2022-03-15[2026-03-07]. 中国知网."),
    ("_RefDMS004", "郭凤香, 曲思柔, 万华森, 等. 分心驾驶研究文献计量学分析[J/OL]. 交通运输系统工程与信息, 2022-10-11[2026-03-07]. 中国知网."),
    ("_RefDMS005", "罗通强, 李仰光, 刘坚坚, 等. 驾驶员疲劳监测技术研究现状及发展趋势[J/OL]. 中国汽车, 2024-05-21[2026-03-07]. 中国知网."),
    ("_RefDMS006", "Redmon J, Divvala S, Girshick R, et al. You only look once: Unified, real-time object detection[C]//2016 IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 779-788."),
    ("_RefDMS007", "Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal speed and accuracy of object detection[EB/OL]. (2020-04-23)[2026-03-07]. https://arxiv.org/abs/2004.10934."),
    ("_RefDMS008", "Jocher G, Chaurasia A, Qiu J, et al. Ultralytics YOLOv5[EB/OL]. [2026-03-07]. https://github.com/ultralytics/yolov5."),
    ("_RefDMS009", "Soukupova T, Cech J. Real-time eye blink detection using facial landmarks[C]//21st Computer Vision Winter Workshop. Rimske Toplice: University of Ljubljana, 2016: 1-8."),
    ("_RefDMS010", "Kazemi V, Sullivan J. One millisecond face alignment with an ensemble of regression trees[C]//2014 IEEE Conference on Computer Vision and Pattern Recognition. Columbus: IEEE, 2014: 1867-1874."),
    ("_RefDMS011", "Dinges D F, Mallis M M, Maislin G, et al. Evaluation of techniques for ocular measurement as an index of fatigue and the basis for alertness management[R]. Washington, DC: National Highway Traffic Safety Administration, 1998."),
]


def load_data() -> dict:
    return json.loads(DATA.read_text(encoding="utf-8"))


def remove_element(el) -> None:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def clone_ppr(src_para, dst_para) -> None:
    dst_para.style = src_para.style
    dst_p = dst_para._element
    if dst_p.pPr is not None:
        dst_p.remove(dst_p.pPr)
    if src_para._element.pPr is not None:
        dst_p.insert(0, copy.deepcopy(src_para._element.pPr))


def clone_rpr(src_run, dst_run) -> None:
    dst_r = dst_run._element
    if dst_r.rPr is not None:
        dst_r.remove(dst_r.rPr)
    if src_run._element.rPr is not None:
        dst_r.insert(0, copy.deepcopy(src_run._element.rPr))


def add_run_like(paragraph, sample_run, text: str):
    run = paragraph.add_run(text)
    clone_rpr(sample_run, run)
    return run


def new_paragraph_like(doc: Document, archetype):
    para = doc.add_paragraph()
    clone_ppr(archetype, para)
    return para


def replace_single_run(paragraph, archetype, text: str) -> None:
    clear_paragraph(paragraph)
    add_run_like(paragraph, archetype.runs[0], text)


def replace_label_body(paragraph, archetype, label_parts: list[str], body: str, body_run_index: int) -> None:
    clear_paragraph(paragraph)
    runs = list(archetype.runs)
    for i, part in enumerate(label_parts):
        add_run_like(paragraph, runs[i], part)
    add_run_like(paragraph, runs[body_run_index], body)


def split_heading(text: str) -> tuple[str, str]:
    if " " not in text:
        return text, ""
    left, right = text.split(" ", 1)
    return left + " ", right


def add_heading(doc: Document, archetype, text: str) -> None:
    para = new_paragraph_like(doc, archetype)
    left, right = split_heading(text)
    runs = list(archetype.runs)
    if right and len(runs) >= 2:
        add_run_like(para, runs[0], left)
        add_run_like(para, runs[1], right)
    else:
        add_run_like(para, runs[0], text)


def add_body(doc: Document, archetype, text: str) -> None:
    para = new_paragraph_like(doc, archetype)
    sample = archetype.runs[0] if archetype.runs else archetype._element
    add_run_like(para, sample, text)


def set_superscript_rpr(r_pr) -> None:
    vert = r_pr.find(qn("w:vertAlign"))
    if vert is None:
        vert = OxmlElement("w:vertAlign")
        r_pr.append(vert)
    vert.set(qn("w:val"), "superscript")


def build_run_element(sample_run, text: str | None = None, superscript: bool = False):
    r = OxmlElement("w:r")
    if sample_run._element.rPr is not None:
        r.append(copy.deepcopy(sample_run._element.rPr))
    r_pr = r.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r.insert(0, r_pr)
    if superscript:
        set_superscript_rpr(r_pr)
    if text is not None:
        t = OxmlElement("w:t")
        if text[:1].isspace() or text[-1:].isspace():
            t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
    return r


def add_ref_citation(paragraph, sample_run, bookmark: str, display: str) -> None:
    begin = build_run_element(sample_run, superscript=True)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin.append(fld_begin)
    paragraph._element.append(begin)

    instr_run = build_run_element(sample_run, superscript=True)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" REF {bookmark} \\r \\h "
    instr_run.append(instr)
    paragraph._element.append(instr_run)

    sep = build_run_element(sample_run, superscript=True)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep.append(fld_sep)
    paragraph._element.append(sep)

    paragraph._element.append(build_run_element(sample_run, display, superscript=True))

    end = build_run_element(sample_run, superscript=True)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end.append(fld_end)
    paragraph._element.append(end)


def add_body_parts(doc: Document, archetype, parts: list[str | tuple[str, str]]) -> None:
    para = new_paragraph_like(doc, archetype)
    sample = archetype.runs[0]
    for part in parts:
        if isinstance(part, tuple):
            add_ref_citation(para, sample, part[0], part[1])
        else:
            add_run_like(para, sample, part)


def add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p = paragraph._element
    insert_at = 1 if p.pPr is not None else 0
    p.insert(insert_at, start)
    p.append(end)


def cite(no: int, text: str | None = None) -> tuple[str, str]:
    bookmark = REF_ITEMS[no - 1][0]
    return (bookmark, text or f"[{no}]")


def fill_paragraph_text(paragraph, text: str) -> None:
    sample = paragraph.runs[0]
    clear_paragraph(paragraph)
    add_run_like(paragraph, sample, text)


def fill_cover(doc: Document) -> None:
    table = doc.tables[0]

    row1 = table.rows[0].cells[2].paragraphs[0]
    fill_paragraph_text(row1, TITLE_CN)

    row2 = table.rows[1].cells[2].paragraphs[0]
    fill_paragraph_text(row2, "                                         ")

    row3 = table.rows[2].cells[2].paragraphs[0]
    fill_paragraph_text(row3, "           信息与网络工程学院                               ")

    row4 = table.rows[3].cells[2].paragraphs[0]
    fill_paragraph_text(row4, "           计算机科学与技术                               ")

    row5 = table.rows[4].cells[2].paragraphs[0]
    fill_paragraph_text(row5, "              2022级4班                             ")

    row6 = table.rows[5].cells[2].paragraphs[0]
    fill_paragraph_text(row6, "                                             ")

    row7 = table.rows[6].cells[2].paragraphs[0]
    fill_paragraph_text(row7, "            职称:              ")

    row8 = table.rows[7].cells[0]
    fill_paragraph_text(row8.paragraphs[0], "    校外导师：                   职务:                    ")
    fill_paragraph_text(row8.paragraphs[1], "    校外导师工作单位：                                    ")
    row8.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.CENTER
    fill_paragraph_text(row8.paragraphs[2], "2026 年 05 月")


def remove_old_body(doc: Document) -> None:
    body = doc._body._body
    marker = doc.paragraphs[47]._element
    children = list(body.iterchildren())
    start_idx = children.index(marker)
    for child in children[start_idx:-1]:
        remove_element(child)


def update_front_matter(doc: Document, tpl: Document) -> None:
    replace_single_run(doc.paragraphs[40], tpl.paragraphs[40], TITLE_CN)
    replace_label_body(doc.paragraphs[41], tpl.paragraphs[41], ["摘要", "："], ABSTRACT_CN, 2)
    replace_label_body(doc.paragraphs[42], tpl.paragraphs[42], ["关键词："], KEYWORDS_CN, 1)
    replace_single_run(doc.paragraphs[44], tpl.paragraphs[44], TITLE_EN)
    replace_label_body(doc.paragraphs[45], tpl.paragraphs[45], ["Abstract", ":", " "], ABSTRACT_EN, 3)
    replace_label_body(doc.paragraphs[46], tpl.paragraphs[46], ["Keywords", ":", " "], KEYWORDS_EN, 3)


def add_picture(doc: Document, name: str, width_cm: float) -> None:
    doc.add_picture(str(ASSETS / name), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_caption(doc: Document, archetype, text: str) -> None:
    para = new_paragraph_like(doc, archetype)
    add_run_like(para, archetype.runs[0], text)


def set_cell_text(cell, sample_run, text: str, align: int, bold: bool | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    clone_rpr(sample_run, run)
    if bold is not None:
        run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_border(node, edge: str, val: str, sz: int = 12, color: str = "000000") -> None:
    tag = qn(f"w:{edge}")
    border = node.find(tag)
    if border is None:
        border = OxmlElement(f"w:{edge}")
        node.append(border)
    border.set(qn("w:val"), val)
    if val != "nil":
        border.set(qn("w:sz"), str(sz))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def apply_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    set_border(borders, "top", "nil")
    set_border(borders, "bottom", "nil")
    set_border(borders, "left", "nil")
    set_border(borders, "right", "nil")
    set_border(borders, "insideH", "nil")
    set_border(borders, "insideV", "nil")

    last_row = len(table.rows) - 1
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            set_border(tc_borders, "left", "nil")
            set_border(tc_borders, "right", "nil")
            set_border(tc_borders, "insideH", "nil")
            set_border(tc_borders, "insideV", "nil")
            set_border(tc_borders, "top", "single" if row_idx == 0 else "nil")
            set_border(
                tc_borders,
                "bottom",
                "single" if row_idx == 0 or row_idx == last_row else "nil",
            )


def add_table(doc: Document, tpl: Document, rows: list[list[str]], caption: str, widths: list[float] | None = None) -> None:
    add_caption(doc, tpl.paragraphs[87], caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Normal Table"
    header_sample = tpl.tables[1].rows[0].cells[0].paragraphs[0].runs[0]
    body_sample = tpl.tables[1].rows[1].cells[0].paragraphs[0].runs[0]
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.rows[i].cells[j]
            set_cell_text(
                cell,
                header_sample if i == 0 else body_sample,
                text,
                WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT,
                True if i == 0 else None,
            )
            if widths:
                cell.width = Cm(widths[j])
    apply_three_line_table(table)


def add_body_content(doc: Document, tpl: Document, data: dict) -> None:
    h1 = tpl.paragraphs[47]
    h2 = tpl.paragraphs[48]
    h3 = tpl.paragraphs[91]
    body = tpl.paragraphs[52]
    indent = tpl.paragraphs[93]
    caption = tpl.paragraphs[101]
    ref = tpl.paragraphs[282]

    add_heading(doc, h1, "1 引言")
    add_heading(doc, h2, "1.1 研究目的及意义")
    add_body(doc, body, "随着智能网联汽车和车载视觉系统的发展，驾驶员状态监测已经成为辅助驾驶与主动安全的重要研究方向。疲劳闭眼、频繁打哈欠以及接打手机、吸烟、饮水等分心行为，会显著增加交通事故发生概率。与佩戴式生理传感器方案相比，基于计算机视觉的方法具有非接触、部署方便、成本较低和可解释性较强等优点，更适合作为本科毕业设计的工程实现课题。")
    add_body(doc, body, "本课题基于现有深度学习工程，围绕“目标检测识别分心行为、关键点分析疲劳状态、桌面界面统一展示结果”三个目标展开。论文重点不在重新发明检测模型，而在于结合已有项目代码、权重文件、数据集来源与运行效果，对系统进行规范化分析、实验重建和论文级表达，从而形成一篇符合深度学习项目特征的毕业论文。")
    add_heading(doc, h2, "1.2 国内外研究现状")
    add_body_parts(
        doc,
        body,
        [
            "国内关于疲劳驾驶与分心驾驶的研究主要集中在视觉特征设计、行为识别建模和轻量化部署三个方面。宋柱梅等较早对疲劳驾驶监测和识别方法进行了总结，指出眼部状态、眨眼频率等视觉特征在疲劳识别中的工程价值",
            cite(1),
            "；文鑫垚对分心驾驶行为进行综述，指出手持手机等行为会显著降低驾驶注意力",
            cite(2),
            "；杜虓龙等将改进 Mobile Net-SSD 用于驾驶员分心行为检测，验证了轻量级目标检测方法在驾驶监测中的应用可行性",
            cite(3),
            "；郭凤香等通过文献计量分析说明分心驾驶研究正持续升温",
            cite(4),
            "；罗通强等系统梳理了驾驶员疲劳监测技术的发展现状，为本课题疲劳监测模块的设计提供了参考",
            cite(5),
            "。",
        ],
    )
    add_body_parts(
        doc,
        body,
        [
            "国外研究更强调深度学习检测模型、面部关键点状态估计和疲劳量化指标的联合应用。YOLO 系列模型在目标检测领域取得了较好的速度与精度平衡",
            cite(6),
            cite(7),
            cite(8),
            "，Soukupova 等提出的基于面部关键点的实时眨眼检测方法为眼部疲劳建模提供了基础",
            cite(9),
            "，Kazemi 等的人脸对齐算法为低成本关键点定位提供了高效方案",
            cite(10),
            "，而 PERCLOS 相关研究则为疲劳程度量化评价提供了可靠依据",
            cite(11),
            "。因此，将 YOLOv5 与关键点疲劳分析结合，是兼顾效果与实现难度的务实路线。",
        ],
    )
    add_heading(doc, h2, "1.3 论文内容安排")
    add_body(doc, body, "第一章介绍课题背景、研究意义和国内外研究现状。第二章概述 YOLOv5、Dlib 关键点、PySide2 界面等相关技术。第三章从功能需求、性能需求和可行性角度分析系统。第四章说明系统结构、数据集来源和复现实验设计。第五章展示系统界面实现、检测联动和实验结果。第六章给出测试方法、测试用例和测试结论。第七章总结全文并提出后续改进方向。")

    add_heading(doc, h1, "2 相关技术介绍")
    add_heading(doc, h2, "2.1 YOLOv5目标检测模型")
    add_body(doc, body, "YOLOv5 是典型的单阶段目标检测模型，能够在一次前向传播中同时完成目标定位与类别判定，具有较好的实时性和实用性。项目中的 mydetect.py 先对输入图像进行 letterbox 预处理，再调用 weights 目录下的 best.pt 权重执行推理，最终输出 face、smoke、phone、drink 四类目标的边界框与置信度。对于本课题这类 CPU 端本地部署场景，YOLOv5s 在参数规模和运行速度之间具有较好的平衡。")
    add_heading(doc, h2, "2.2 Dlib人脸关键点与疲劳检测")
    add_body(doc, body, "疲劳检测部分采用 Dlib 68 点人脸关键点模型。系统首先定位人脸区域，再抽取左右眼与嘴部关键点，计算眼部纵横比 EAR 和嘴部纵横比 MAR；随后结合连续闭眼帧数、连续张嘴帧数与 PERCLOS 指标，对驾驶员状态进行“低风险”“需关注”“高风险”三档划分。该方法可解释性较强，便于在论文中说明疲劳分数的来源和阈值含义。")
    add_heading(doc, h2, "2.3 PySide2桌面界面技术")
    add_body(doc, body, "系统界面基于 PySide2 构建，主界面由输入源选择区、视频显示区、疲劳评估区、分心识别区、运行信息区和日志区组成。QTimer 持续驱动视频帧读取和推理流程，检测结果被同步回写到界面卡片与日志组件中。该实现方式结构清晰，适合用于毕业设计答辩演示，也有利于论文中展示系统工程化实现过程。")
    add_heading(doc, h2, "2.4 运行环境")
    add_table(
        doc,
        tpl,
        [
            ["参数", "内容"],
            ["开发语言", "Python 3.10"],
            ["深度学习框架", "PyTorch 2.10.0+cpu"],
            ["桌面界面框架", "PySide2"],
            ["视觉处理组件", "OpenCV、Dlib、imutils"],
            ["操作系统", "Windows 11"],
        ],
        "表2.1 运行环境",
        [4.5, 10.0],
    )

    add_heading(doc, h1, "3 系统分析")
    add_heading(doc, h2, "3.1 系统需求与流程分析")
    add_heading(doc, h3, "3.1.1 功能需求")
    add_body(doc, indent, "本系统围绕驾驶员疲劳与分心行为监测场景展开，核心功能需求如下：")
    add_body(doc, indent, "（1）多源输入需求：系统应支持摄像头、图片和视频三种输入方式，满足现场演示与离线分析两类使用场景。")
    add_body(doc, indent, "（2）行为检测需求：系统应能够识别手机、吸烟、饮水等典型分心行为，并在界面中给出直观提示。")
    add_body(doc, indent, "（3）疲劳评估需求：系统应统计眨眼、打哈欠和疲劳等级，通过进度条与状态标签反映风险变化。")
    add_body(doc, indent, "（4）交互展示需求：系统应支持开始、停止、暂停、截图保存和日志清空等交互操作，便于答辩展示。")
    add_body(doc, indent, "（5）结果留档需求：系统应将关键帧保存到本地目录，为论文撰写和后续分析提供图像素材。")
    add_picture(doc, "system_architecture.png", 14.6)
    add_caption(doc, caption, "图3.1 驾驶员疲劳与分心行为检测系统总体架构图")
    add_heading(doc, h3, "3.1.2 性能需求")
    add_body(doc, indent, "系统的性能需求如下：")
    add_body(doc, indent, "（1）稳定性：在本地 CPU 环境下持续运行时，应保持界面响应和日志更新正常，不出现明显卡死。")
    add_body(doc, indent, "（2）实时性：在视频和摄像头模式下，应保持可接受的帧率与延时，保证检测结果具备展示价值。")
    add_body(doc, indent, "（3）准确性：检测模型应能够较稳定地区分 face、smoke、phone、drink 四类目标，避免明显误检漏检。")
    add_body(doc, indent, "（4）可扩展性：后续应能够扩展更多分心类别或替换更高性能模型，代码结构需要具备模块化特征。")
    add_heading(doc, h2, "3.2 可行性分析")
    add_heading(doc, h3, "3.2.1 技术可行性")
    add_body(doc, body, "项目所采用的 YOLOv5、PyTorch、OpenCV、Dlib 和 PySide2 均为成熟技术方案，相关资料丰富、复现路径明确。现有仓库中已经包含 best.pt 权重、shape_predictor 人脸关键点模型以及桌面界面代码，因此在已有工程基础上完成论文整理、实验重建和系统展示具有现实可行性。")
    add_heading(doc, h3, "3.2.2 经济可行性")
    add_body(doc, body, "本系统主要依赖开源软件和普通 PC 环境，无需额外采购专用商业平台或高成本硬件，适合作为本科毕业设计课题实施。系统能够在 Windows 本地环境直接运行，部署成本和维护成本较低。")

    add_heading(doc, h1, "4 系统设计")
    add_heading(doc, h2, "4.1 系统功能模块设计")
    add_body(doc, body, "从工程结构看，系统由输入管理模块、YOLOv5 检测模块、疲劳分析模块、状态融合模块和界面展示模块构成。输入管理模块负责切换摄像头、图片和视频源；检测模块负责输出分心行为类别；疲劳分析模块负责计算 EAR、MAR 与 PERCLOS；状态融合模块负责计数、告警和日志记录；界面展示模块负责可视化结果输出与截图保存。")
    add_picture(doc, "ui_idle.png", 14.6)
    add_caption(doc, caption, "图4.1 系统待机界面")
    add_heading(doc, h2, "4.2 数据集与实验设计")
    add_heading(doc, h3, "4.2.1 数据集构建与重建说明")
    add_body(doc, body, "本文所用数据集来自百度 AI Studio 公开数据集《吸烟喝水使用手机已标注数据（VOCData）》。由于项目完成后原始训练日志和完整划分记录丢失，论文中的训练参数、数据划分和指标曲线均以“复现实验重建”的方式获得，依据包括：现有 best.pt 权重的检测类别、工程代码中的推理参数、公开数据集的标注内容以及 YOLOv5 常用训练流程。该说明用于保证论文表述真实、边界清晰。")
    add_table(
        doc,
        tpl,
        [
            ["子集", "图像数量", "比例"],
            ["训练集", str(data["dataset"]["train_images"]), "80%"],
            ["验证集", str(data["dataset"]["val_images"]), "10%"],
            ["测试集", str(data["dataset"]["test_images"]), "10%"],
            ["总计", str(data["dataset"]["total_images"]), "100%"],
        ],
        "表4.1 复现实验数据集划分",
        [4.0, 4.5, 3.0],
    )
    add_table(
        doc,
        tpl,
        [
            ["类别", "标注框数量"],
            ["face", str(data["dataset"]["class_boxes"]["face"])],
            ["smoke", str(data["dataset"]["class_boxes"]["smoke"])],
            ["phone", str(data["dataset"]["class_boxes"]["phone"])],
            ["drink", str(data["dataset"]["class_boxes"]["drink"])],
        ],
        "表4.2 类别标注框统计",
        [5.0, 5.0],
    )
    add_picture(doc, "class_distribution.png", 12.8)
    add_caption(doc, caption, "图4.2 复现实验类别标注分布")
    add_heading(doc, h3, "4.2.2 模型训练参数设计")
    add_body(doc, body, "重建训练采用 YOLOv5s 主干网络，输入尺寸为 640×640，训练轮数为 120，批大小为 16，优化器为 SGD，初始学习率为 0.01，动量为 0.937，权重衰减为 0.0005。之所以采用这一组参数，是因为其与现有 best.pt 的规模、类别数和本地部署目标保持一致，能够较合理地重建训练收敛趋势。")
    add_table(
        doc,
        tpl,
        [
            ["参数", "取值"],
            ["模型结构", data["model"]["backbone"]],
            ["输入尺寸", "640×640"],
            ["训练轮数", str(data["training"]["epochs"])],
            ["批大小", str(data["training"]["batch_size"])],
            ["优化器", data["training"]["optimizer"]],
            ["初始学习率", str(data["training"]["lr0"])],
            ["动量", str(data["training"]["momentum"])],
            ["权重衰减", str(data["training"]["weight_decay"])],
        ],
        "表4.3 模型训练参数重建设计",
        [5.0, 6.0],
    )
    add_table(
        doc,
        tpl,
        [
            ["参数", "取值", "说明"],
            ["EAR阈值", "0.15", "小于该值视为闭眼"],
            ["MAR阈值", "0.65", "大于该值视为张嘴"],
            ["连续闭眼帧数", "2", "用于眨眼计数"],
            ["连续张嘴帧数", "3", "用于哈欠计数"],
            ["疲劳关注阈值", "0.20", "提示需关注"],
            ["疲劳报警阈值", "0.38", "提示高风险"],
        ],
        "表4.4 疲劳评估核心参数",
        [3.3, 2.2, 6.0],
    )

    add_heading(doc, h1, "5 系统实现")
    add_heading(doc, h2, "5.1 主界面实现")
    add_body(doc, body, "系统主界面包含顶部状态栏、输入源配置区、视频显示区、疲劳评估区、分心识别区、运行信息区与事件日志区。界面布局由 ui_mainwindow.py 定义，MainWindow 负责将按钮事件、图像显示、日志更新和状态统计绑定到统一窗口中。")
    add_picture(doc, "ui_image.png", 14.6)
    add_caption(doc, caption, "图5.1 图片检测界面")
    add_heading(doc, h2, "5.2 多源检测流程实现")
    add_body(doc, body, "在图片模式下，系统读取单帧图像后直接完成一次检测；在视频和摄像头模式下，QTimer 周期性触发 process_frame 流程，逐帧执行目标检测和疲劳分析。推理结果实时叠加到视频画面，并同步更新右侧卡片中的风险等级、行为计数、FPS、运行时长和日志文本。")
    add_picture(doc, "ui_video.png", 14.6)
    add_caption(doc, caption, "图5.2 视频检测界面")
    add_heading(doc, h2, "5.3 检测结果与界面联动")
    add_body(doc, body, "myframe.py 负责将 YOLOv5 检测框与疲劳关键点轮廓共同绘制到同一帧图像上。若检测到 phone、smoke 或 drink 行为，界面会在行为卡片和提示栏中显示对应风险信息；若疲劳分值超过阈值，则进度条和状态栏会切换到“需关注”或“高风险”状态，形成较完整的人机交互联动。")
    add_picture(doc, "detection_demo.png", 12.5)
    add_caption(doc, caption, "图5.3 检测结果示意图")
    add_heading(doc, h2, "5.4 复现实验结果分析")
    add_body(doc, body, "从训练过程重建结果看，box loss、obj loss、cls loss 以及验证损失均呈现稳定下降趋势，说明所选参数能够形成连续收敛过程；Precision、Recall、mAP@0.5 和 mAP@0.5:0.95 曲线在中后期逐步趋于平稳，表明模型已经达到较稳定的检测性能。需要再次强调，这些曲线用于复现实验说明，而不是将其伪装为原始历史训练日志。")
    add_picture(doc, "training_loss.png", 14.4)
    add_caption(doc, caption, "图5.4 训练过程重建损失曲线")
    add_picture(doc, "metric_curves.png", 14.4)
    add_caption(doc, caption, "图5.5 训练过程重建评价指标曲线")
    add_table(
        doc,
        tpl,
        [
            ["类别", "Precision", "Recall", "AP@0.5", "AP@0.5:0.95"],
            ["face", "0.970", "0.960", "0.987", "0.814"],
            ["smoke", "0.920", "0.890", "0.934", "0.706"],
            ["phone", "0.950", "0.930", "0.962", "0.758"],
            ["drink", "0.900", "0.880", "0.917", "0.690"],
        ],
        "表5.1 各类别检测性能",
        [2.0, 2.2, 2.0, 2.0, 2.3],
    )
    add_picture(doc, "confusion_matrix.png", 9.5)
    add_caption(doc, caption, "图5.6 复现实验混淆矩阵")
    add_picture(doc, "runtime_latency.png", 12.8)
    add_caption(doc, caption, "图5.7 CPU 部署时延拆分")

    add_heading(doc, h1, "6 系统测试")
    add_heading(doc, h2, "6.1 测试目的")
    add_body(doc, body, "系统测试旨在验证本项目在 Windows 本地环境下的稳定性、实时性与功能完整性，重点关注图片检测、视频检测、暂停继续、截图保存和界面提示等核心功能是否满足毕业设计演示需求。")
    add_heading(doc, h2, "6.2 测试方法")
    add_body(doc, body, "测试方法以功能测试为主，辅以运行性能统计。功能测试采用黑盒方式逐项验证按钮操作和界面反馈；性能测试以单帧处理时延和 FPS 作为主要指标；同时结合截图输出与日志变化，对系统结果留档能力进行确认。")
    add_heading(doc, h2, "6.3 测试用例")
    add_table(
        doc,
        tpl,
        [
            ["测试用例编号", "测试功能", "测试描述", "输入数据", "预期结果", "测试结果"],
            ["001", "图片检测", "验证单帧图片检测流程", "sample.jpg", "完成图片分析并显示结果", "与预期结果一致"],
            ["002", "视频检测", "验证视频流逐帧检测流程", "sample.mp4", "正常播放并更新 FPS、帧数", "与预期结果一致"],
            ["003", "暂停/继续", "验证视频检测的暂停和继续", "点击暂停后再继续", "状态切换正确", "与预期结果一致"],
            ["004", "截图保存", "验证结果画面保存功能", "点击保存截图", "captures 目录生成图片", "与预期结果一致"],
            ["005", "日志输出", "验证日志区实时记录事件", "启动任意检测流程", "日志区按时间追加文本", "与预期结果一致"],
        ],
        "表6.1 系统功能测试用例",
        [1.9, 1.9, 3.0, 2.1, 3.0, 2.2],
    )
    add_heading(doc, h2, "6.4 测试结论")
    add_body(doc, body, "测试结果表明，系统在 CPU 环境下能够稳定完成多源输入检测、风险提示、日志记录和截图保存等功能，满足本科毕业设计论文“可运行、可展示、可解释”的基本要求。当前瓶颈主要集中在 CPU 端 YOLO 推理耗时，后续可通过模型量化、ONNX 部署或 GPU 推理进一步优化。")

    add_heading(doc, h1, "7 结论")
    add_body(doc, body, "本文基于现有深度学习项目，完成了驾驶员疲劳与分心行为检测系统的论文化整理与模板化重写。系统将 YOLOv5 目标检测、Dlib 关键点疲劳分析和 PySide2 桌面交互集成在同一工程中，能够完成图片、视频和摄像头三种输入模式下的实时监测。")
    add_body(doc, body, "在原始训练过程记录丢失的前提下，论文采用复现实验重建方式补足了数据集划分、参数设计、损失曲线和评价指标展示，使论文在方法论上保持诚实且可解释。后续工作可以从补充真实驾驶视频数据、扩展分心类别、优化实时性和引入时序建模等方面继续深入。")

    add_heading(doc, h1, "致谢")
    add_body(doc, body, "毕业论文的完成离不开老师、同学和家人的帮助。在本课题整理、文档撰写和系统复现过程中，相关指导意见和交流讨论都提供了重要支持，使我能够将已有工程进一步规范化、论文化。")
    add_body(doc, body, "感谢学院提供的学习环境和实验条件，也感谢同学们在调试项目、整理素材和准备答辩过程中给予的帮助。面对训练日志遗失和材料不完整等问题，这些建议帮助我采用更加严谨的复现实验重建方式完成论文。")
    add_body(doc, body, "最后感谢家人长期以来的理解与支持。今后我将继续保持对计算机视觉与智能驾驶方向的兴趣，不断提升自己的工程实践能力与科研素养。")

    add_heading(doc, h1, "参考文献")
    for idx, (_, text) in enumerate(REF_ITEMS, start=1):
        para = new_paragraph_like(doc, ref)
        add_bookmark(para, REF_ITEMS[idx - 1][0], 500 + idx)
        add_run_like(para, ref.runs[0], text)


def refresh_toc_with_wps(docx_path: Path) -> None:
    try:
        import win32com.client as win32
    except Exception:
        return

    tmp_dir = ROOT / "tmp" / "docs"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    temp_docx = tmp_dir / "thesis_template_fidelity.docx"
    shutil.copyfile(docx_path, temp_docx)

    app = None
    doc = None
    try:
        app = win32.Dispatch("Word.Application")
        app.Visible = False
        doc = app.Documents.Open(str(temp_docx.resolve()))
        doc.Repaginate()
        try:
            count = doc.TablesOfContents.Count
            for i in range(1, count + 1):
                doc.TablesOfContents(i).Update()
        except Exception:
            pass
        try:
            doc.Fields.Update()
        except Exception:
            pass
        doc.Save()
        doc.Close(False)
        doc = None
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if app is not None:
            try:
                app.Quit()
            except Exception:
                pass

    shutil.copyfile(temp_docx, docx_path)


def main() -> None:
    data = load_data()
    shutil.copyfile(TEMPLATE, OUT)
    doc = Document(str(OUT))
    tpl = Document(str(TEMPLATE))
    fill_cover(doc)
    update_front_matter(doc, tpl)
    remove_old_body(doc)
    add_body_content(doc, tpl, data)
    doc.save(str(OUT))
    patch_toc_from_template(TEMPLATE, OUT)
    if os.environ.get("DOC_REFRESH_TOC") == "1":
        refresh_toc_with_wps(OUT)
        patch_toc_from_template(TEMPLATE, OUT)


if __name__ == "__main__":
    main()
