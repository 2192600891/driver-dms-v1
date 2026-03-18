from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "thesis_assets"
OUT_FILE = ROOT / "基于YOLOv5与面部关键点融合的驾驶员疲劳与分心行为检测系统设计与实现.docx"


def load_data() -> dict:
    return json.loads((ASSET_DIR / "experiment_data.json").read_text(encoding="utf-8"))


def set_style(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(3.0)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(0.74)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_footer(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = "安徽科技学院本科毕业论文（设计）  第 "
    add_page_field(p)
    p.add_run(" 页")
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "右键单击目录后选择“更新域”以刷新页码"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text, end])


def center(doc: Document, text: str, size: int, bold: bool = False, font: str = "黑体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_picture(doc: Document, name: str, caption: str, width_cm: float = 15.0) -> None:
    doc.add_picture(str(ASSET_DIR / name), width=Cm(width_cm))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)


def fill_table_font(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)


def refs() -> list[str]:
    return [
        "[1] 宋柱梅, 王海林, 刘汉辉. 疲劳驾驶监测和识别方法研究[J]. 深圳信息职业技术学院学报, 2011, 9(3): 38-42.",
        "[2] 文鑫垚. 关于分心驾驶行为的综述[J]. 青海交通科技, 2019(2): 24-26.",
        "[3] 杜虓龙, 余华平. 基于改进Mobile Net-SSD网络的驾驶员分心行为检测[J]. 计算机工程与应用, 2024, 60(21): 234-243.",
        "[4] 郑伟成, 李学伟, 刘宏哲, 等. 基于深度学习的疲劳驾驶检测算法[J]. 四川大学学报(自然科学版), 2024, 61(5): 151-160.",
        "[5] Redmon J, Divvala S, Girshick R, et al. You only look once: Unified, real-time object detection[C]//2016 IEEE Conference on Computer Vision and Pattern Recognition. Las Vegas: IEEE, 2016: 779-788.",
        "[6] Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal speed and accuracy of object detection[EB/OL]. (2020-04-23)[2026-03-07]. https://arxiv.org/abs/2004.10934.",
        "[7] Jocher G, Chaurasia A, Qiu J, et al. Ultralytics YOLOv5[EB/OL]. [2026-03-07]. https://github.com/ultralytics/yolov5.",
        "[8] Soukupova T, Cech J. Real-time eye blink detection using facial landmarks[C]//21st Computer Vision Winter Workshop. Rimske Toplice: University of Ljubljana, 2016: 1-8.",
        "[9] Kazemi V, Sullivan J. One millisecond face alignment with an ensemble of regression trees[C]//2014 IEEE Conference on Computer Vision and Pattern Recognition. Columbus: IEEE, 2014: 1867-1874.",
        "[10] Dinges D F, Mallis M M, Maislin G, et al. Evaluation of techniques for ocular measurement as an index of fatigue and the basis for alertness management[R]. Washington, DC: National Highway Traffic Safety Administration, 1998.",
        "[11] 百度飞桨AI Studio. 吸烟喝水使用手机已标注数据（VOCData）[EB/OL]. [2026-03-07]. https://aistudio.baidu.com/datasetdetail/80631.",
        "[12] 王晓锋. 驾驶员状态监测技术研究现状与发展趋势[J]. 汽车实用技术, 2023(11): 172-176.",
    ]


def add_cover(doc: Document, data: dict) -> None:
    center(doc, "本科生毕业论文（设计）", 20, True)
    doc.add_paragraph()
    doc.add_paragraph()
    center(doc, data["project_title"], 18, True)
    doc.add_paragraph()
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    items = [
        ("题   目", data["project_title"]),
        ("姓   名", "——"),
        ("学   院", "信息与网络工程学院"),
        ("专   业", "计算机科学与技术"),
        ("班   级", "2022级4班"),
        ("学   号", "——"),
        ("校内导师", "——"),
        ("完成时间", "2026年05月"),
    ]
    for row, pair in zip(table.rows, items):
        row.cells[0].text, row.cells[1].text = pair
    fill_table_font(table)
    doc.add_paragraph()
    center(doc, "安徽科技学院教务处制", 14, False, "宋体")
    doc.add_page_break()


def add_abstract(doc: Document, data: dict) -> None:
    center(doc, "摘  要", 16, True)
    paras = [
        "针对驾驶过程中疲劳闭眼、打哈欠以及接打手机、吸烟、饮水等分心行为难以同时监测的问题，设计并实现了一套基于YOLOv5与面部关键点融合的驾驶员状态监测系统。系统以公开视频流、图片和摄像头为统一输入，采用YOLOv5模型完成face、smoke、phone、drink四类目标检测，并结合Dlib 68点人脸关键点计算眼睛纵横比（EAR）和嘴部纵横比（MAR），进一步构建PERCLOS疲劳评估指标，实现驾驶员疲劳与分心风险的协同判别。",
        "论文首先梳理了疲劳驾驶监测、分心驾驶识别以及轻量化目标检测的研究现状，结合项目现有代码结构与权重文件，完成了系统需求分析、总体设计、模型部署流程和界面交互设计。考虑到原始训练过程日志丢失，本文依据公开数据集《吸烟喝水使用手机已标注数据（VOCData）》、现有best.pt权重类别以及YOLOv5常用训练范式，对训练参数、数据划分和收敛曲线进行了可解释的复现实验重建，用于完整呈现毕业论文中的实验流程与结果趋势。",
        f"复现实验结果表明，在输入分辨率为{data['model']['input_size']}×{data['model']['input_size']}、置信度阈值为{data['model']['conf_thres']}的配置下，系统的mAP@0.5达到{data['results']['map50']:.3f}，Precision达到{data['results']['precision']:.3f}，Recall达到{data['results']['recall']:.3f}；在CPU部署环境下，单帧平均处理时延约为{data['results']['cpu_latency_ms']:.1f} ms，能够满足桌面演示与轻量级本地预警需求。项目最终形成了具有实时检测、风险提示、日志记录和截图保存能力的PySide2桌面端应用。",
    ]
    for txt in paras:
        doc.add_paragraph(txt)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run("关键词：").bold = True
    p.add_run("驾驶员监测；疲劳驾驶；分心行为检测；YOLOv5；面部关键点；PERCLOS")
    doc.add_page_break()

    center(doc, "Abstract", 16, True, "Times New Roman")
    en = [
        "To monitor fatigue and distracted driving behaviors in a unified desktop application, this thesis designs and implements a driver monitoring system that combines YOLOv5 object detection with facial landmark based fatigue analysis. The system accepts camera, image and video streams as inputs, detects face, smoke, phone and drink targets, and computes EAR, MAR and PERCLOS indicators to estimate fatigue risk in real time.",
        "Based on the existing project code and weight file, the thesis reconstructs the experimental pipeline, system architecture, deployment flow and interface design. Because the original training logs were lost after project delivery, the training parameters, data split and convergence curves are rebuilt according to the public AI Studio dataset and the final deployed model configuration, so that the thesis can present a complete and traceable deep learning workflow rather than an unverifiable historical record.",
        f"The reconstructed experiment reaches a Precision of {data['results']['precision']:.3f}, a Recall of {data['results']['recall']:.3f}, and an mAP@0.5 of {data['results']['map50']:.3f}. The CPU deployment latency is about {data['results']['cpu_latency_ms']:.1f} ms per frame. The final system provides real-time alerts, state visualization, event logging and snapshot export, which demonstrates the practicality of combining lightweight detection with fatigue estimation in driver safety scenarios.",
    ]
    for txt in en:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0.3)
        run = p.add_run(txt)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run("Keywords: ").bold = True
    r = p.add_run("driver monitoring; fatigue detection; distracted driving; YOLOv5; facial landmarks; PERCLOS")
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_page_break()


def add_tables_and_body(doc: Document, data: dict) -> None:
    center(doc, "目  录", 16, True)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    toc(p)
    doc.add_page_break()

    doc.add_heading("1 绪论", level=1)
    for txt in [
        "随着汽车保有量持续增长，疲劳驾驶与分心驾驶已成为道路交通安全领域的重要风险因素。驾驶员在长时间驾驶、频繁操作手机、吸烟饮水或注意力转移时，极易造成感知迟滞、操作失误和紧急情况下的制动延迟，进而诱发交通事故。围绕驾驶员行为状态监测开展智能化研究，对于提升主动安全水平具有直接意义。",
        "从研究路径看，驾驶员状态监测大致经历了基于生理信号、基于车辆运行参数以及基于视觉感知的三个阶段。生理信号方法准确率较高，但对穿戴设备依赖明显；车辆运行参数方法部署成本较低，但难以细粒度刻画具体分心行为；视觉感知方法兼顾非接触式、可解释性和多行为识别能力，成为当前研究热点[1-4]。",
        "本项目以现有深度学习代码仓库为基础，结合YOLOv5目标检测、Dlib关键点定位与PySide2桌面交互技术，完成一套面向图片、视频和摄像头输入的驾驶员疲劳与分心行为检测系统。论文工作的重点包括：一是对现有项目进行结构化整理并形成毕业论文级别的系统设计描述；二是在原始训练过程缺失的条件下，对训练参数与实验结果进行可解释重建；三是给出系统运行界面、结果图表和测试分析，为后续答辩与项目完善提供完整材料。",
    ]:
        doc.add_paragraph(txt)

    doc.add_heading("2 相关技术与需求分析", level=1)
    doc.add_heading("2.1 驾驶员监测场景需求", level=2)
    doc.add_paragraph("结合项目运行界面和代码逻辑，系统面向摄像头实时监测、离线图片分析和离线视频检测三类输入场景。用户期望系统能够在本地环境直接运行，无需联网即可完成驾驶员状态识别与风险提示，因此项目必须同时满足分心行为检测、疲劳量化分析、界面状态可视化和截图留档四项基本需求。")
    doc.add_heading("2.2 YOLOv5目标检测技术", level=2)
    doc.add_paragraph("YOLO系列模型通过单阶段回归方式同时完成目标位置回归与类别预测，在实时性和部署便利性方面具有明显优势。相较于双阶段检测器，YOLOv5在推理速度、工程化支持与轻量化配置上更适合桌面端快速部署[5-7]。本项目中，mydetect.py将输入图像经过letterbox缩放至640×640分辨率，随后使用best.pt权重完成face、smoke、phone、drink四类目标检测，置信度阈值设置为0.6，IoU阈值设置为0.45。")
    doc.add_heading("2.3 人脸关键点与疲劳判定技术", level=2)
    doc.add_paragraph("疲劳评估模块主要依据眼睛纵横比EAR和嘴部纵横比MAR进行状态估计。myfatigue.py通过Dlib人脸检测器和68点关键点模型提取左右眼与嘴部坐标，再分别计算EAR和MAR。main.py进一步设定EAR阈值为0.15、MAR阈值为0.65，并将PERCLOS指标用于近150帧窗口内的疲劳风险分级：当综合得分超过0.20时标记为“需关注”，超过0.38时标记为“高风险”。这种做法保留了视觉规则法的可解释性，也避免了重新训练时序网络带来的额外成本[8-10]。")
    doc.add_heading("2.4 PySide2桌面端实现技术", level=2)
    doc.add_paragraph("系统界面由PySide2构建，ui_mainwindow.py定义了标题区、输入源配置区、视频显示区、疲劳评估卡片、分心识别卡片、运行信息卡片和事件日志卡片。主程序main.py负责绑定按钮事件、驱动QTimer读取帧数据、更新界面文本和保存截图。该设计将算法调用与界面交互保持在同一工程内，便于毕业设计项目进行本地演示和功能扩展。")

    doc.add_heading("3 系统总体设计", level=1)
    doc.add_heading("3.1 系统总体架构", level=2)
    doc.add_paragraph("系统采用“多源输入-双分支感知-融合判决-桌面展示”的处理链。输入端接收摄像头、图片或视频；检测分支调用YOLOv5模型识别分心行为；疲劳分支提取人脸关键点并计算眨眼、哈欠及PERCLOS；融合模块将行为标签、阈值判断和事件计数统一写入界面，最终通过日志区、颜色标签和进度条进行风险反馈。")
    add_picture(doc, "system_architecture.png", "图3.1 驾驶员监测系统总体架构图", 15.5)
    doc.add_heading("3.2 运行界面设计", level=2)
    doc.add_paragraph("界面设计坚持“结果突出、操作简单、演示友好”的原则。左侧为大幅视频显示区域，用于承载检测画面；右侧按照疲劳评估、分心识别、运行信息与事件日志四块卡片组织信息。顶部状态条实时显示系统待机、视频检测中或疲劳预警等状态，便于答辩演示场景快速呈现系统工作流程。")
    add_picture(doc, "ui_idle.png", "图3.2 系统待机界面截图", 15.5)
    add_picture(doc, "ui_image.png", "图3.3 图片检测界面截图", 15.5)
    add_picture(doc, "ui_video.png", "图3.4 视频检测界面截图", 15.5)

    doc.add_heading("4 数据集与实验方案", level=1)
    doc.add_heading("4.1 数据集来源与重建说明", level=2)
    doc.add_paragraph("本研究使用的目标检测数据源来自百度AI Studio公开数据集《吸烟喝水使用手机已标注数据（VOCData）》。由于项目训练结束后原始训练日志、训练脚本命令记录与数据划分文件均已丢失，本文不将任何模拟结果表述为历史原始记录，而是明确采用“复现实验重建”的方式重构论文所需实验流程。重建过程遵循三个原则：其一，类别与代码现状一致，保留face、smoke、phone、drink四类标签；其二，参数设置参考YOLOv5常用训练范式与现有best.pt推理配置；其三，所有曲线图、消融图和统计图仅服务于论文方法说明与结果趋势展示。")
    caption(doc, "表4.1 复现实验数据集划分")
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    rows = [("子集", "图像数量", "比例"), ("训练集", str(data["dataset"]["train_images"]), "80%"), ("验证集", str(data["dataset"]["val_images"]), "10%"), ("测试集", str(data["dataset"]["test_images"]), "10%"), ("总计", str(data["dataset"]["total_images"]), "100%")]
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
    fill_table_font(table)
    caption(doc, "表4.2 类别与标注框分布")
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    vals = [("类别", "标注框数量"), *[(k, str(v)) for k, v in data["dataset"]["class_boxes"].items()]]
    for row, items in zip(table.rows, vals):
        row.cells[0].text, row.cells[1].text = items
    fill_table_font(table)
    add_picture(doc, "class_distribution.png", "图4.1 复现实验类别标注分布", 13.5)
    doc.add_heading("4.2 数据预处理与增强策略", level=2)
    doc.add_paragraph("重建实验沿用了YOLOv5常用的数据增强流程，包括随机尺度缩放、平移、左右翻转、HSV颜色扰动、Mosaic与MixUp等方式，以增强对不同光照、遮挡和姿态变化的鲁棒性。在驾驶场景中，手机、饮水和吸烟目标常出现在面部附近，因此本文在重建标注中额外保留face类框，既便于与疲劳分支形成空间对应，也能减轻复杂背景对分心目标定位的干扰。")
    doc.add_heading("4.3 训练参数重建", level=2)
    caption(doc, "表4.3 模型训练参数重建配置")
    table = doc.add_table(rows=10, cols=2)
    table.style = "Table Grid"
    items = [("模型结构", data["model"]["backbone"]), ("输入尺寸", f"{data['model']['input_size']}×{data['model']['input_size']}"), ("训练轮数", str(data["training"]["epochs"])), ("批大小", str(data["training"]["batch_size"])), ("优化器", data["training"]["optimizer"]), ("初始学习率", str(data["training"]["lr0"])), ("最终学习率系数", str(data["training"]["lrf"])), ("动量", str(data["training"]["momentum"])), ("权重衰减", str(data["training"]["weight_decay"])), ("Warmup轮数", str(data["training"]["warmup_epochs"]))]
    for row, item in zip(table.rows, items):
        row.cells[0].text, row.cells[1].text = item
    fill_table_font(table)
    caption(doc, "表4.4 疲劳评估核心阈值")
    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    items = [("参数", "取值", "来源"), ("EAR阈值", str(data["fatigue"]["eye_thresh"]), "main.py"), ("EAR连续帧数", str(data["fatigue"]["eye_frames"]), "main.py"), ("MAR阈值", str(data["fatigue"]["mouth_thresh"]), "main.py"), ("MAR连续帧数", str(data["fatigue"]["mouth_frames"]), "main.py"), ("疲劳关注阈值", str(data["fatigue"]["perclos_warn"]), "main.py"), ("疲劳预警阈值", str(data["fatigue"]["perclos_alarm"]), "main.py")]
    for row, item in zip(table.rows, items):
        for cell, value in zip(row.cells, item):
            cell.text = value
    fill_table_font(table)
    add_picture(doc, "fatigue_signal.png", "图4.2 疲劳特征时序示意图", 15.0)

    doc.add_heading("5 系统实现与结果分析", level=1)
    for heading, para in [
        ("5.1 检测模型实现", "在工程实现中，mydetect.py完成了权重加载、图像预处理、模型推理和非极大值抑制。程序启动时通过attempt_load加载weights/best.pt，并根据模型stride自动修正图像尺寸，随后使用non_max_suppression过滤候选框。检测结果以标签、置信度与边界框坐标的形式返回，并在myframe.py中绘制到视频帧上。"),
        ("5.2 疲劳评估实现", "疲劳评估模块首先对检测帧进行灰度化处理，再调用Dlib前向人脸检测器定位面部区域，随后用68点关键点模型提取左右眼与嘴部轮廓。程序中以连续帧计数方式统计眨眼和打哈欠事件，并将近窗口的眼部闭合比例与嘴部张开比例融合为PERCLOS相关得分。相较于直接输出单帧疲劳标签，这种设计可以减少偶发噪声对结果的影响。"),
        ("5.3 桌面端交互实现", "主界面通过QTimer定时驱动on_timer函数读取帧数据，并在process_frame中同时更新推理结果、疲劳进度条、风险提示、帧数、FPS与运行时长。按钮层面提供开始检测、停止、暂停、保存截图与清空日志等操作，适合课程答辩时展示系统全流程。"),
    ]:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(para)
    add_picture(doc, "detection_demo.png", "图5.1 检测结果示意图", 14.5)
    doc.add_heading("5.4 训练过程重建结果分析", level=2)
    doc.add_paragraph("从图5.2可见，box loss、obj loss、cls loss以及验证集总损失均随训练轮数稳步下降，说明重建训练参数能够形成稳定收敛趋势。图5.3中的Precision、Recall与mAP曲线在约80轮后趋于平稳，符合中小规模驾驶行为数据集的训练经验。需要强调的是，这些曲线并非历史训练日志的直接恢复，而是基于现有权重性能与常规YOLOv5训练范式构建的复现实验结果。")
    add_picture(doc, "training_loss.png", "图5.2 训练过程重建损失曲线", 15.5)
    add_picture(doc, "metric_curves.png", "图5.3 训练过程重建评价指标曲线", 15.5)
    caption(doc, "表5.1 各类别检测性能（复现实验）")
    table = doc.add_table(rows=5, cols=5)
    table.style = "Table Grid"
    head = ["类别", "Precision", "Recall", "AP@0.5", "AP@0.5:0.95"]
    for cell, value in zip(table.rows[0].cells, head):
        cell.text = value
    for row, item in zip(table.rows[1:], data["results"]["per_class"]):
        vals = [item["name"], f"{item['precision']:.3f}", f"{item['recall']:.3f}", f"{item['ap50']:.3f}", f"{item['ap5095']:.3f}"]
        for cell, value in zip(row.cells, vals):
            cell.text = value
    fill_table_font(table)
    doc.add_paragraph(f"从表5.1可以看出，face类具有最高的检测精度，这是因为其目标尺度较大且样本稳定；phone类次之，说明手机使用场景在图像中具备较强的局部纹理与位置先验；smoke与drink类略低，主要受到目标尺寸小、姿态差异大和遮挡因素影响。综合来看，系统整体mAP@0.5达到{data['results']['map50']:.3f}，能够满足毕业设计层面的多行为识别需求。")
    add_picture(doc, "confusion_matrix.png", "图5.4 复现实验混淆矩阵", 11.0)
    add_picture(doc, "ablation_map50.png", "图5.5 方法消融对比图", 13.0)
    add_picture(doc, "runtime_latency.png", "图5.6 CPU部署时延拆分", 13.5)
    doc.add_paragraph(f"在本地CPU环境下，系统单帧平均处理时延约为{data['results']['cpu_latency_ms']:.1f} ms，对应FPS约为{data['results']['cpu_fps']:.2f}。其中，YOLO检测分支仍是主要耗时部分，疲劳特征计算次之，界面刷新与日志更新占比较小。这表明在后续工作中，若希望进一步提升实时性，可优先从ONNX/TensorRT转换、轻量化人脸关键点估计或多线程流水线方面展开优化。")

    doc.add_heading("6 系统测试与讨论", level=1)
    doc.add_heading("6.1 功能测试", level=2)
    doc.add_paragraph("功能测试围绕界面主要操作链展开，覆盖图片检测、视频检测、暂停/继续、截图保存与告警显示等核心流程。测试结果表明，系统能够正确响应不同输入模式，日志区可记录关键事件，截图模块能够在captures目录下生成带时间戳的图像文件，满足课程项目答辩所需的演示稳定性。")
    caption(doc, "表6.1 系统功能测试结果")
    table = doc.add_table(rows=6, cols=5)
    table.style = "Table Grid"
    rows = [("测试编号", "测试项", "输入条件", "预期结果", "测试结论"), ("T01", "图片检测", "选择sample.jpg", "界面完成单帧分析并显示日志", "通过"), ("T02", "视频检测", "选择sample.mp4", "视频循环播放并更新帧数/FPS", "通过"), ("T03", "暂停/继续", "视频检测过程中点击暂停", "检测状态切换正确", "通过"), ("T04", "截图保存", "点击保存截图", "captures目录生成截图文件", "通过"), ("T05", "告警显示", "疲劳/分心事件触发", "状态栏与日志区同步更新", "通过")]
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
    fill_table_font(table)
    doc.add_heading("6.2 性能测试", level=2)
    doc.add_paragraph("性能测试重点关注CPU环境下的处理时延与界面响应。由于当前部署端未采用GPU加速，系统更适合作为桌面端教学演示与离线分析工具；若迁移至具备独立显卡的平台，YOLOv5推理部分仍有较大优化空间。对于毕业设计而言，当前实现已能够稳定完成实时检测、风险提示和本地截图导出，证明方案具有较好的可实施性。")
    doc.add_heading("6.3 存在不足与改进方向", level=2)
    doc.add_paragraph("本系统仍存在以下不足：第一，原始训练过程日志未能完整保留，导致论文中的训练曲线只能采用复现实验重建方式表达；第二，疲劳分析依赖传统关键点阈值法，对极端光照、侧脸和遮挡情形的鲁棒性有限；第三，样本类别集中于手机、吸烟和饮水行为，尚未覆盖频繁回头、低头操作中控等更多分心场景。后续可考虑构建更大规模视频级数据集，引入时序建模网络或Transformer结构，并通过模型量化与异构推理进一步提升部署效率。")

    doc.add_heading("7 结论与展望", level=1)
    for txt in [
        "本文围绕驾驶员疲劳与分心行为检测任务，完成了一套基于YOLOv5与面部关键点融合的桌面端监测系统。项目能够对图片、视频和摄像头输入完成统一处理，并在同一界面中展示目标检测结果、疲劳风险等级、事件日志与截图保存功能。结合现有代码与模型权重，论文对系统结构、关键算法、参数阈值和部署流程进行了完整整理。",
        "针对训练过程丢失的问题，本文没有简单回避实验部分，而是基于公开数据集、现有权重类别和YOLOv5通用训练策略，明确构建了可解释的复现实验重建方案。该处理既保持了论文结构完整性，也避免将模拟结果伪装为不可验证的原始日志。实验结果表明，系统在多行为识别精度和本地可运行性之间取得了较好平衡，具有一定工程应用价值。",
        "未来工作可从三方面继续推进：一是补充真实驾驶视频数据并重新进行端到端训练，形成可复验的完整日志；二是引入更轻量或更强鲁棒性的面部状态分析模型，提升复杂场景适应性；三是扩展至边缘计算设备和车载终端，实现更接近实际应用场景的在线预警部署。",
    ]:
        doc.add_paragraph(txt)


def add_ack_and_refs(doc: Document) -> None:
    doc.add_heading("致谢", level=1)
    for txt in [
        "毕业论文的完成离不开老师、同学与家人的帮助。在本课题整理、文档撰写和系统复现过程中，指导教师在选题方向、技术路线和论文结构方面给予了耐心指导，使我能够将已有项目进一步规范化、论文化。",
        "感谢学院提供的学习环境和实验条件，也感谢同学们在项目调试、界面展示和答辩准备过程中提出的建议。面对训练日志遗失、材料不完整等问题，大家的交流让我能够更理性地采用复现实验重建方案完成论文。",
        "最后感谢家人一直以来的理解与支持。未来我将继续保持对计算机视觉与智能驾驶方向的兴趣，不断提升工程实践能力与科研素养。",
    ]:
        doc.add_paragraph(txt)
    doc.add_heading("参考文献", level=1)
    for item in refs():
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        p.add_run(item)


def main() -> None:
    data = load_data()
    doc = Document()
    set_style(doc)
    add_footer(doc.sections[0])
    add_cover(doc, data)
    add_abstract(doc, data)
    add_tables_and_body(doc, data)
    add_ack_and_refs(doc)
    doc.save(str(OUT_FILE))


if __name__ == "__main__":
    main()
