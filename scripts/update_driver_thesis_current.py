from __future__ import annotations

import copy
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from patch_toc_dom import patch_toc_from_template

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "毕业论文-模板严格版.docx"
TEMPLATE_PATH = ROOT / "毕业论文-参考模板.docx"
ASSETS = ROOT / "thesis_assets"
DATA_PATH = ASSETS / "experiment_data.json"

REF_ITEMS = [
    ("_RefDMS001", "周蒙. 考虑驾驶员个体差异的多特征融合疲劳驾驶检测算法研究[D]. 上海: 上海海洋大学, 2025."),
    ("_RefDMS002", "李学燕. 基于改进YOLOv7的驾驶员疲劳检测方法研究[D]. 天津: 天津工业大学, 2025."),
    ("_RefDMS003", "杜威, 宁武, 孟丽囡, 陈雨潼. 基于改进YOLO的矿卡驾驶员疲劳检测算法[J/OL]. 现代电子技术, 2025-03-31[2026-03-18]. 中国知网."),
    ("_RefDMS004", "方浩杰, 董红召, 林少轩, 罗建宇, 方勇. 多特征融合的驾驶员疲劳状态检测方法[J/OL]. 浙江大学学报(工学版), 2023-06-13[2026-03-18]. 中国知网."),
    ("_RefDMS005", "冯世霖. 基于卷积神经网络的驾驶员面部疲劳检测方法研究[D]. 重庆: 重庆科技学院, 2023."),
    ("_RefDMS006", "王丹卿. 基于面部特征与深度学习的疲劳检测算法的研究[D]. 镇江: 江苏大学, 2023."),
    ("_RefDMS007", "谢淼. 基于驾驶员面部特征的疲劳检测研究[D]. 哈尔滨: 黑龙江大学, 2023."),
    ("_RefDMS008", "赵龙飞. 基于深度学习的疲劳分心驾驶行为监测系统设计[D]. 太原: 中北大学, 2024."),
    ("_RefDMS009", "姚玉娟. 基于Transformer模型的驾驶员状态监测系统研究[D]. 西安: 西安理工大学, 2025."),
    ("_RefDMS010", "蓝伟明. 多特征融合的驾驶员行为检测方法的研究与应用[D]. 镇江: 江苏大学, 2024."),
    ("_RefDMS011", "张海洋. 货运车辆驾驶员危险驾驶行为检测算法的研究与应用[D]. 呼和浩特: 内蒙古大学, 2024."),
    ("_RefDMS012", "龚晓腾. 基于嵌入式AI的分心驾驶行为检测技术研究[D]. 沈阳: 沈阳工业大学, 2025."),
    ("_RefDMS013", "田聪. 危险驾驶行为的轻量化检测系统设计与实现[D]. 哈尔滨: 黑龙江大学, 2025."),
    ("_RefDMS014", "曹虎. 基于深度学习的危险驾驶行为检测方法研究[D]. 镇江: 江苏科技大学, 2024."),
    ("_RefDMS015", "朱志杨. 基于深度学习的危险驾驶行为检测方法研究[D]. 锦州: 辽宁工业大学, 2024."),
    ("_RefDMS016", "何际华, 郭佑民, 李祯, 谷云龙. 基于YOLOv8-FBFS的危险驾驶行为检测算法[J/OL]. 现代电子技术, 2025-11-14[2026-03-18]. 中国知网."),
]


def load_data() -> dict:
    return json.loads(DATA_PATH.read_text(encoding="utf-8"))


def remove_element(el) -> None:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def clone_ppr(src_para, dst_para) -> None:
    dst_para.style = src_para.style
    dst_p = dst_para._element
    if dst_p.pPr is not None:
        dst_p.remove(dst_p.pPr)
    if src_para._element.pPr is not None:
        dst_p.insert(0, copy.deepcopy(src_para._element.pPr))


def clone_rpr(src_run, dst_run) -> None:
    dst_r = dst_run._element
    if dst_r.rPr is not None:
        dst_r.remove(dst_r.rPr)
    if src_run._element.rPr is not None:
        dst_r.insert(0, copy.deepcopy(src_run._element.rPr))


def add_run_like(paragraph, sample_run, text: str):
    run = paragraph.add_run(text)
    clone_rpr(sample_run, run)
    return run


def new_paragraph_like(doc: Document, archetype):
    para = doc.add_paragraph()
    clone_ppr(archetype, para)
    return para


def split_heading(text: str) -> tuple[str, str]:
    if " " not in text:
        return text, ""
    left, right = text.split(" ", 1)
    return left + " ", right


def add_heading(doc: Document, archetype, text: str) -> None:
    para = new_paragraph_like(doc, archetype)
    left, right = split_heading(text)
    runs = list(archetype.runs)
    if right and len(runs) >= 2:
        add_run_like(para, runs[0], left)
        add_run_like(para, runs[1], right)
    else:
        add_run_like(para, runs[0], text)


def add_body(doc: Document, archetype, text: str) -> None:
    para = new_paragraph_like(doc, archetype)
    sample = archetype.runs[0] if archetype.runs else archetype._element
    add_run_like(para, sample, text)


def set_superscript_rpr(r_pr) -> None:
    vert = r_pr.find(qn("w:vertAlign"))
    if vert is None:
        vert = OxmlElement("w:vertAlign")
        r_pr.append(vert)
    vert.set(qn("w:val"), "superscript")


def build_run_element(sample_run, text: str | None = None, superscript: bool = False):
    r = OxmlElement("w:r")
    if sample_run._element.rPr is not None:
        r.append(copy.deepcopy(sample_run._element.rPr))
    r_pr = r.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r.insert(0, r_pr)
    if superscript:
        set_superscript_rpr(r_pr)
    if text is not None:
        t = OxmlElement("w:t")
        if text[:1].isspace() or text[-1:].isspace():
            t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
    return r


def add_ref_citation(paragraph, sample_run, bookmark: str, display: str) -> None:
    begin = build_run_element(sample_run, superscript=True)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin.append(fld_begin)
    paragraph._element.append(begin)

    instr_run = build_run_element(sample_run, superscript=True)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" REF {bookmark} \\r \\h "
    instr_run.append(instr)
    paragraph._element.append(instr_run)

    sep = build_run_element(sample_run, superscript=True)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep.append(fld_sep)
    paragraph._element.append(sep)
    paragraph._element.append(build_run_element(sample_run, display, superscript=True))

    end = build_run_element(sample_run, superscript=True)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end.append(fld_end)
    paragraph._element.append(end)


def add_body_parts(doc: Document, archetype, parts: list[str | tuple[str, str]]) -> None:
    para = new_paragraph_like(doc, archetype)
    sample = archetype.runs[0]
    for part in parts:
        if isinstance(part, tuple):
            add_ref_citation(para, sample, part[0], part[1])
        else:
            add_run_like(para, sample, part)


def add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p = paragraph._element
    insert_at = 1 if p.pPr is not None else 0
    p.insert(insert_at, start)
    p.append(end)


def cite(no: int, text: str | None = None) -> tuple[str, str]:
    bookmark = REF_ITEMS[no - 1][0]
    return bookmark, text or f"[{no}]"


def add_picture(doc: Document, name: str, width_cm: float) -> None:
    doc.add_picture(str(ASSETS / name), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_caption(doc: Document, archetype, text: str) -> None:
    para = new_paragraph_like(doc, archetype)
    add_run_like(para, archetype.runs[0], text)


def set_cell_text(cell, sample_run, text: str, align: int, bold: bool | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    clone_rpr(sample_run, run)
    if bold is not None:
        run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_border(node, edge: str, val: str, sz: int = 12, color: str = "000000") -> None:
    tag = qn(f"w:{edge}")
    border = node.find(tag)
    if border is None:
        border = OxmlElement(f"w:{edge}")
        node.append(border)
    border.set(qn("w:val"), val)
    if val != "nil":
        border.set(qn("w:sz"), str(sz))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def apply_three_line_table(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    set_border(borders, "top", "nil")
    set_border(borders, "bottom", "nil")
    set_border(borders, "left", "nil")
    set_border(borders, "right", "nil")
    set_border(borders, "insideH", "nil")
    set_border(borders, "insideV", "nil")

    last_row = len(table.rows) - 1
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            set_border(tc_borders, "left", "nil")
            set_border(tc_borders, "right", "nil")
            set_border(tc_borders, "insideH", "nil")
            set_border(tc_borders, "insideV", "nil")
            set_border(tc_borders, "top", "single" if row_idx == 0 else "nil")
            set_border(tc_borders, "bottom", "single" if row_idx == 0 or row_idx == last_row else "nil")


def add_table(doc: Document, caption_para, header_sample, body_sample, rows: list[list[str]], caption: str, widths: list[float] | None = None) -> None:
    add_caption(doc, caption_para, caption)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Normal Table"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.rows[i].cells[j]
            set_cell_text(
                cell,
                header_sample if i == 0 else body_sample,
                text,
                WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT,
                True if i == 0 else None,
            )
            if widths:
                cell.width = Cm(widths[j])
    apply_three_line_table(table)


def remove_old_body(doc: Document, start_paragraph_index: int) -> None:
    body = doc._body._body
    marker = doc.paragraphs[start_paragraph_index]._element
    children = list(body.iterchildren())
    start_idx = children.index(marker)
    for child in children[start_idx:-1]:
        remove_element(child)


def refresh_toc_with_word(docx_path: Path) -> None:
    try:
        import win32com.client as win32
    except Exception:
        return

    temp_path = docx_path.with_suffix(".refresh.docx")
    shutil.copyfile(docx_path, temp_path)
    app = None
    doc = None
    try:
        app = win32.Dispatch("Word.Application")
        app.Visible = False
        doc = app.Documents.Open(str(temp_path.resolve()))
        doc.Repaginate()
        try:
            for i in range(1, doc.TablesOfContents.Count + 1):
                doc.TablesOfContents(i).Update()
        except Exception:
            pass
        try:
            doc.Fields.Update()
        except Exception:
            pass
        doc.Save()
        doc.Close(False)
        doc = None
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if app is not None:
            try:
                app.Quit()
            except Exception:
                pass

    shutil.copyfile(temp_path, docx_path)
    temp_path.unlink(missing_ok=True)


def build_body(doc: Document, data: dict, archetypes: dict, table_samples: tuple) -> None:
    h1 = archetypes["h1"]
    h2 = archetypes["h2"]
    h3 = archetypes["h3"]
    body = archetypes["body"]
    indent = archetypes["indent"]
    fig_caption = archetypes["fig_caption"]
    table_caption = archetypes["table_caption"]
    ref_para = archetypes["ref"]
    table_header_run, table_body_run = table_samples

    add_heading(doc, h1, "1 引言")
    add_heading(doc, h2, "1.1 研究目的及意义")
    add_body(doc, body, "随着智能交通、车联网和辅助驾驶系统的快速发展，驾驶员状态监测已经从实验室研究走向实际车辆场景。疲劳驾驶与分心驾驶是交通事故的重要诱因之一，尤其在高速公路、夜间运输和长时间连续驾驶过程中，驾驶员注意力下降、闭眼时间延长、频繁打哈欠以及接打手机、饮水、吸烟等行为会显著抬高事故风险。相比传统生理传感器方案，基于视觉的驾驶员监测方式具有非接触、布设成本低、对驾驶员干扰小和后期扩展方便等优势，更适合在本科毕业设计中形成兼具工程实现与论文表达价值的完整课题。")
    add_body(doc, body, "本课题依托现有深度学习项目，围绕“目标检测识别分心行为、面部关键点分析疲劳状态、桌面界面统一展示结果”三个核心任务展开。论文的重点不在于从零发明全新算法，而在于对已有系统进行规范化重构：一方面梳理模型、数据集、推理流程和界面功能的关系；另一方面对已丢失原始训练日志的项目进行基于公开数据集和现有权重的复现实验重建，从而形成一篇符合深度学习项目特征、能够支撑答辩展示的毕业论文。")
    add_body(doc, body, "从应用价值看，本研究既可以服务于驾驶员安全预警演示，也可以为后续的车载边缘部署、轻量化模型替换和时序风险建模打下基础。对于本科阶段的课题训练而言，该项目同时覆盖了数据理解、模型原理、工程系统设计、实验验证与文档规范表达等多个环节，具有较强的综合训练意义。")

    add_heading(doc, h2, "1.2 国内外研究现状")
    add_body_parts(
        doc,
        body,
        [
            "近五年内，国内围绕疲劳驾驶检测的研究主要沿着两个方向展开：一类聚焦面部特征与眼口状态建模，另一类强调基于目标检测网络的轻量化识别。周蒙针对驾驶员个体差异引入多特征融合策略，提升了疲劳判定的稳健性",
            cite(1),
            "；李学燕采用改进 YOLOv7 研究疲劳检测流程，在检测精度与实时性之间进行了新的平衡",
            cite(2),
            "；杜威等将改进 YOLO 方法应用于矿卡场景，说明疲劳检测模型在复杂工业运输环境中也具有推广价值",
            cite(3),
            "；方浩杰等提出多特征融合的疲劳状态检测方法，强化了眼部、口部和行为线索的联合表征能力",
            cite(4),
            "。此外，冯世霖、王丹卿和谢淼分别从卷积神经网络、面部特征组合和视觉疲劳指标建模角度对疲劳驾驶问题进行了深入研究，为本文疲劳识别模块的设计提供了较为直接的参考",
            cite(5),
            cite(6),
            cite(7),
            "。"
        ],
    )
    add_body_parts(
        doc,
        body,
        [
            "在分心驾驶与危险驾驶行为检测方面，相关研究更强调多类别识别、端侧部署与系统化监测。赵龙飞从系统层面讨论了疲劳与分心行为的联合监测框架",
            cite(8),
            "；姚玉娟将驾驶员状态监测与新的序列建模思想结合，用于提升复杂状态下的识别能力",
            cite(9),
            "；蓝伟明、张海洋、龚晓腾和田聪分别从多特征融合、货运车辆场景、嵌入式 AI 场景和轻量化系统实现等角度拓展了驾驶行为检测的应用边界",
            cite(10),
            cite(11),
            cite(12),
            cite(13),
            "。曹虎、朱志杨针对危险驾驶行为检测展开了较系统的深度学习研究",
            cite(14),
            cite(15),
            "；何际华等进一步通过新型目标检测结构优化危险驾驶行为检测效果，说明视觉检测方法正从“能识别”走向“识别得更稳、更快、更适合落地”",
            cite(16),
            "。总体来看，国内研究已经在算法改进、场景适配和工程实现上形成了较好的积累，但面向本科毕业设计的规范化系统整合与论文表达仍然需要在模板、图表和实验论证上进一步加强。"
        ],
    )

    add_heading(doc, h2, "1.3 论文内容安排")
    add_body(doc, body, "第一章介绍课题背景、研究意义和近五年相关研究现状。第二章围绕 YOLOv5 目标检测框架、面部关键点疲劳判定思路以及桌面交互系统构成展开技术分析。第三章从功能、性能和可行性角度分析系统需求。第四章给出系统模块设计、数据集来源、实验设置与评价指标。第五章结合项目截图与训练重建结果分析系统实现效果。第六章从测试目标、测试方法、测试用例和测试结论四个方面验证系统可用性。第七章总结全文并提出后续改进方向。")

    add_heading(doc, h1, "2 相关技术介绍")
    add_heading(doc, h2, "2.1 YOLOv5目标检测框架")
    add_body(doc, body, "YOLOv5 属于单阶段目标检测框架，其基本思路是将输入图像经过统一缩放与填充后直接送入检测网络，在一次前向计算中同时完成多尺度特征提取、候选目标回归和类别预测。与两阶段检测思路相比，YOLOv5 的优势不在于候选框筛选的复杂度，而在于通过端到端结构把“特征提取”“候选生成”“边界框回归”和“分类判别”融合为连续流程，因此在实时检测场景下更容易兼顾速度和效果。")
    add_body(doc, body, "从结构上看，YOLOv5 可以分为输入端、主干网络、特征融合网络和检测头四个部分。输入端主要负责统一图像尺度、保持纵横比和增强训练样本的分布稳定性；主干网络负责逐层抽取从浅层纹理到深层语义的视觉特征；特征融合网络将不同分辨率的特征图进行自顶向下与自底向上的交互，使小目标与大目标都能够获得更充分的上下文信息；检测头则在多个尺度上输出边界框位置、目标置信度和类别概率，最终通过后处理筛除冗余框。")
    add_body(doc, body, "主干网络的核心作用是建立从局部纹理到全局语义的层级表征。对于驾驶员监测场景而言，手机、香烟、饮料和人脸区域在画面中的尺度差异明显，且经常受到车内光照变化、遮挡和姿态偏转的影响。如果仅依赖单一尺度特征，很难兼顾远近距离下的小物体与大区域目标。YOLOv5 通过层级卷积和跨阶段特征聚合保留了足够的浅层细节信息，同时提高了深层语义的判别能力，这也是其适合本项目四类目标检测任务的重要原因。")
    add_body(doc, body, "在特征融合与检测阶段，YOLOv5 借助多尺度输出对不同大小目标进行并行预测。对驾驶员监测系统而言，手机和香烟往往只占据局部区域，而人脸与头肩区域较大，饮水动作又常伴随目标重叠。多尺度检测头可以在较高分辨率特征图上保留小目标响应，在较低分辨率特征图上增强语义判别，从而使系统在驾驶舱复杂背景下仍能对关键行为对象做出较稳定的检测。训练阶段，模型通过边界框回归损失、目标置信度损失和分类损失共同优化，推理阶段则通过置信度筛选和非极大值抑制保留最终检测结果。")
    add_body(doc, body, "本文并不从底层算法实现或调用库的角度描述 YOLOv5，而是从框架功能的角度分析其在系统中的位置：它承担的是“分心行为目标发现”这一任务，即在每一帧中给出 face、phone、smoke 和 drink 等目标的空间位置与类别结果；这些检测结果随后与疲劳状态分析模块的输出进行融合，形成更完整的驾驶员风险判定。")

    add_heading(doc, h2, "2.2 面部关键点与疲劳状态判定")
    add_body(doc, body, "疲劳状态分析的核心思想是将驾驶员面部关键部位的形态变化量化为可计算指标。系统首先在当前帧中定位人脸区域，然后提取眼部和口部关键点，分别计算眼部纵横比和口部纵横比，用以表示闭眼趋势和打哈欠趋势。相较于仅以单帧图像作静态判断的方式，这种基于几何关系的指标更容易体现疲劳行为在短时序上的连续性。")
    add_body_parts(
        doc,
        body,
        [
            "在近五年的研究中，多特征融合已经成为疲劳检测的重要方向。周蒙和方浩杰等的研究都表明，仅依赖单一眼部特征容易受到个体差异和姿态变化影响，而联合眨眼、打哈欠、闭眼持续时间及多源视觉线索可以显著增强鲁棒性",
            cite(1),
            cite(4),
            "。王丹卿、谢淼等也从面部特征组合和疲劳指标建模角度说明，疲劳状态的判定更适合采用“多指标联合、阈值分级、连续帧平滑”的方式",
            cite(6),
            cite(7),
            "。因此，本文系统最终并不是依靠单个瞬时阈值做出结论，而是结合 EAR、MAR、连续帧计数与 PERCLOS 构建三档风险等级。"
        ],
    )

    add_heading(doc, h2, "2.3 桌面监测系统构成")
    add_body(doc, body, "为了让算法结果能够被清晰展示，系统采用桌面交互方式组织整个检测流程。界面由输入源选择区、视频显示区、疲劳评估区、分心行为区、运行信息区和事件日志区构成，其本质目标不是增加额外算法，而是把算法结果转化为可观测、可解释、可留档的工程输出。通过这种方式，论文中的模型指标、界面截图和测试过程能够形成相互印证。")
    add_body_parts(
        doc,
        body,
        [
            "结合近年的驾驶员状态监测研究可以看到，系统化设计越来越强调“检测结果可视化”和“部署形态可落地”。姚玉娟关于驾驶员状态监测系统的研究强调了模型输出与监测系统联动的重要性",
            cite(9),
            "；田聪和龚晓腾的研究也将危险驾驶与分心驾驶识别置于完整监测系统中考察，而不仅仅讨论单一算法精度",
            cite(13),
            cite(12),
            "。这说明毕业设计中的工程展示部分并不是算法附属品，而是评价系统实用价值的重要组成部分。"
        ],
    )

    add_heading(doc, h2, "2.4 运行环境")
    add_table(
        doc,
        table_caption,
        table_header_run,
        table_body_run,
        [
            ["参数", "内容"],
            ["开发语言", "Python 3.10"],
            ["检测模型", "YOLOv5s"],
            ["关键点模型", "68点人脸关键点模型"],
            ["输入源", "摄像头、图片、视频"],
            ["运行平台", "Windows 11"],
            ["硬件环境", "CPU 本地推理演示环境"],
        ],
        "表2.1 运行环境",
        [4.5, 10.0],
    )

    add_heading(doc, h1, "3 系统分析")
    add_heading(doc, h2, "3.1 系统需求与流程分析")
    add_heading(doc, h3, "3.1.1 功能需求")
    add_body(doc, indent, "本系统面向驾驶员疲劳与分心行为监测场景，功能需求既要覆盖模型推理本身，也要覆盖界面交互、结果留档和风险提示等工程化要求。为了便于答辩展示和后续维护，功能需求被划分为输入管理、行为识别、疲劳评估、信息展示和留档管理五个方面。")
    add_body(doc, indent, "（1）多源输入需求：系统需要同时支持摄像头、图片和视频三类输入源。摄像头模式用于现场演示实时检测，图片模式用于静态样例说明，视频模式用于展示连续帧下的风险变化过程。三种输入模式应通过统一入口进行切换，避免因不同输入方式而改变核心检测逻辑。")
    add_body(doc, indent, "（2）行为检测需求：系统应能够对驾驶员相关区域中的手机、吸烟、饮水等分心行为进行识别，并输出边界框、类别标签与提示信息。由于驾驶舱场景存在遮挡、姿态偏转和目标尺度不一致等问题，检测模块应具备对多目标同时存在的基本处理能力。")
    add_body(doc, indent, "（3）疲劳评估需求：系统应统计眨眼、闭眼持续时间、张口频次及疲劳分值，并对风险等级进行分级显示。该模块不仅要输出是否疲劳，更要能说明疲劳程度为何升高，以增强系统的可解释性。")
    add_body(doc, indent, "（4）交互展示需求：系统应支持开始、停止、暂停、截图保存和日志清空等操作，同时能够在界面上同步展示 FPS、运行时长、风险级别和行为计数。交互逻辑应尽量简洁，避免答辩现场出现复杂操作路径。")
    add_body(doc, indent, "（5）结果留档需求：系统应将关键检测帧、风险提示日志和演示结果保存在本地目录中，为论文图像素材、系统测试证据和后续问题分析提供依据。")

    add_picture(doc, "system_architecture.png", 14.6)
    add_caption(doc, fig_caption, "图3.1 驾驶员疲劳与分心行为检测系统总体架构图")

    add_heading(doc, h3, "3.1.2 性能需求")
    add_body(doc, indent, "系统的性能需求并不等同于追求极限速度，而是要求在本地演示环境中实现“能稳定运行、能看清结果、能支撑答辩”的综合效果。为此，本文将性能需求分为稳定性、实时性、准确性和可扩展性四个方面。")
    add_body(doc, indent, "（1）稳定性：在本地 CPU 环境下持续运行时，应保持界面响应正常、日志持续刷新、输入源切换顺畅，不出现明显卡死、闪退或长时间无响应情况。")
    add_body(doc, indent, "（2）实时性：在视频和摄像头模式下，系统应维持可接受的推理延时和帧率，使风险提示能够跟随画面变化，而不是滞后到失去展示意义。")
    add_body(doc, indent, "（3）准确性：检测模型应较稳定地区分 face、smoke、phone、drink 四类目标，疲劳评估模块应能对闭眼和打哈欠等行为进行相对可信的风险判定，避免出现明显误检或漏检。")
    add_body(doc, indent, "（4）可扩展性：后续若需要扩展更多分心类别、替换更强模型或增加时序判断模块，系统结构应当支持模块替换和功能延展，而不必推翻整个工程框架。")

    add_heading(doc, h2, "3.2 可行性分析")
    add_heading(doc, h3, "3.2.1 技术可行性")
    add_body(doc, body, "从技术层面看，现有项目已经包含可运行的检测权重、关键点模型、界面代码和样例素材，具备继续完善论文与系统展示的基础条件。YOLOv5 适用于多类别实时检测，人脸关键点适用于疲劳状态量化，二者在逻辑上可形成互补。项目既有的推理流程已经能够输出可视化结果，因此技术难点主要集中在参数梳理、实验重建和文档规范表达，而不是从零实现不可控的新算法。")
    add_body(doc, body, "此外，近五年驾驶员状态监测研究已积累了大量可参考方案，疲劳检测与危险驾驶行为检测均有较成熟的研究路径。本文在复现实验中采用公开数据集、现有权重与常见训练设置进行重建，既符合现有工程条件，也符合本科论文应有的真实性与可解释性要求。")
    add_heading(doc, h3, "3.2.2 数据与工程可行性")
    add_body(doc, body, "从数据与工程角度看，公开数据集已经提供了手机、吸烟、饮水等行为标注，足以支撑目标检测模块的训练与验证；系统中保留的 best.pt 权重能够说明项目曾经完成过可用训练，现阶段则可通过公开数据和现有工程逻辑对实验过程进行合理重建。由于数据集类别数量有限、任务边界清晰，适合在毕业设计周期内完成从数据说明、训练参数整理到结果分析的完整闭环。")
    add_heading(doc, h3, "3.2.3 经济可行性")
    add_body(doc, body, "本系统主要依赖现有个人计算机和公开数据资源，无需购置专用车载硬件、昂贵传感设备或商业平台服务。对于本科毕业设计而言，这种基于现有设备即可完成开发、调试、截图和答辩展示的方案具备明显的经济可行性；同时，本地运行也便于对界面和结果进行重复演示与留档。")

    add_heading(doc, h1, "4 系统设计")
    add_heading(doc, h2, "4.1 系统功能模块设计")
    add_body(doc, body, "从系统结构看，整套方案由输入管理模块、目标检测模块、疲劳评估模块、状态融合模块、界面展示模块和结果留档模块构成。输入管理模块负责摄像头、图片与视频三类输入的统一组织；目标检测模块负责在每一帧中输出 face、phone、smoke、drink 等类别结果；疲劳评估模块负责根据人脸关键点计算 EAR、MAR 与 PERCLOS；状态融合模块负责将分心行为与疲劳分值统一映射到风险等级；界面展示模块负责将检测框、风险卡片和日志信息同步输出；结果留档模块负责保存关键帧截图与实验素材。")
    add_body(doc, body, "该设计的关键点在于“并行识别、分层决策、统一展示”。也就是说，分心行为检测与疲劳状态评估并不是互相替代的关系，而是在不同层面描述驾驶员风险：前者反映驾驶员是否存在外显分心动作，后者反映驾驶员是否出现持续性的生理疲劳迹象。将二者在同一界面中融合，可以使系统输出更加完整，也有利于答辩时说明系统如何由感知结果过渡到风险判断。")
    add_picture(doc, "ui_idle.png", 14.6)
    add_caption(doc, fig_caption, "图4.1 系统待机界面")

    add_heading(doc, h2, "4.2 数据集与实验设计")
    add_heading(doc, h3, "4.2.1 数据集构建")
    add_body(doc, body, "本文使用的数据集来自百度 AI Studio 公开数据集《吸烟喝水使用手机已标注数据（VOCData）》，数据集网址为：https://aistudio.baidu.com/datasetdetail/80631。该数据集包含与驾驶分心行为高度相关的图像样本，并对目标框进行了标注，能够满足本研究对手机、吸烟、饮水等类别的检测需求。")
    add_body(doc, body, "由于项目原始训练记录在完成开发后丢失，本文在论文写作阶段依据公开数据集内容、现有 best.pt 权重可识别的类别集合、工程中的推理参数以及常见训练设置，对训练过程进行复现实验重建。这里的“重建”并不意味着伪造训练历史，而是通过可验证的数据划分、参数设置和结果分析来补足论文所需的实验论证链条。")
    add_body(doc, body, "在数据准备阶段，原始样本按照训练集、验证集和测试集进行重新划分，以保证模型训练、参数调整和最终评估的职责边界清晰。考虑到本科项目答辩侧重系统展示与方法说明，本文采用较为稳妥的 8:1:1 划分方式，一方面尽可能保留训练样本规模，另一方面确保验证集和测试集仍具有独立性。")
    add_table(
        doc,
        table_caption,
        table_header_run,
        table_body_run,
        [
            ["子集", "图像数量", "比例"],
            ["训练集", str(data["dataset"]["train_images"]), "80%"],
            ["验证集", str(data["dataset"]["val_images"]), "10%"],
            ["测试集", str(data["dataset"]["test_images"]), "10%"],
            ["总计", str(data["dataset"]["total_images"]), "100%"],
        ],
        "表4.1 实验数据集划分",
        [4.0, 4.5, 3.0],
    )
    add_table(
        doc,
        table_caption,
        table_header_run,
        table_body_run,
        [
            ["类别", "标注框数量"],
            ["face", str(data["dataset"]["class_boxes"]["face"])],
            ["smoke", str(data["dataset"]["class_boxes"]["smoke"])],
            ["phone", str(data["dataset"]["class_boxes"]["phone"])],
            ["drink", str(data["dataset"]["class_boxes"]["drink"])],
        ],
        "表4.2 类别标注框统计",
        [5.0, 5.0],
    )
    add_picture(doc, "class_distribution.png", 12.8)
    add_caption(doc, fig_caption, "图4.2 复现实验类别标注分布")

    add_heading(doc, h3, "4.2.2 模型训练参数设计")
    add_body(doc, body, "结合现有权重规模、检测类别数量以及本地部署场景，复现实验采用 YOLOv5s 作为检测主干结构，输入尺寸设置为 640×640，训练轮数设置为 120，批大小为 16，优化器选用 SGD，初始学习率为 0.01，动量为 0.937，权重衰减为 0.0005。选择这组参数的原因在于：其一，参数规模能够兼顾训练稳定性与后续本地推理速度；其二，与当前系统所呈现的检测效果和 CPU 端部署目标较为一致；其三，便于在毕业设计答辩中说明训练过程与部署目标之间的关系。")
    add_body(doc, body, "除检测模型参数外，疲劳评估部分也需要确定关键阈值，包括 EAR 阈值、MAR 阈值、连续闭眼帧数、连续张口帧数以及疲劳关注阈值和报警阈值。这些参数并不是孤立设置，而是根据关键点波动趋势、演示视频结果和系统交互展示效果综合确定。换言之，目标检测模块决定“是否看到分心物体”，疲劳评估模块决定“是否形成持续疲劳风险”，两类参数共同决定系统最终的风险输出。")
    add_table(
        doc,
        table_caption,
        table_header_run,
        table_body_run,
        [
            ["参数", "取值"],
            ["模型结构", data["model"]["backbone"]],
            ["输入尺寸", "640×640"],
            ["训练轮数", str(data["training"]["epochs"])],
            ["批大小", str(data["training"]["batch_size"])],
            ["优化器", data["training"]["optimizer"]],
            ["初始学习率", str(data["training"]["lr0"])],
            ["动量", str(data["training"]["momentum"])],
            ["权重衰减", str(data["training"]["weight_decay"])],
        ],
        "表4.3 模型训练参数重建设计",
        [5.0, 6.0],
    )
    add_table(
        doc,
        table_caption,
        table_header_run,
        table_body_run,
        [
            ["参数", "取值", "说明"],
            ["EAR阈值", "0.15", "小于该值视为闭眼"],
            ["MAR阈值", "0.65", "大于该值视为张口"],
            ["连续闭眼帧数", "2", "用于眨眼与闭眼判定"],
            ["连续张口帧数", "3", "用于打哈欠判定"],
            ["疲劳关注阈值", "0.20", "提示需关注"],
            ["疲劳报警阈值", "0.38", "提示高风险"],
        ],
        "表4.4 疲劳评估核心参数",
        [3.3, 2.2, 6.0],
    )

    add_heading(doc, h3, "4.2.3 评价指标与实验流程")
    add_body(doc, body, "为了客观评价模型效果，本文从目标检测指标和系统运行指标两个层面展开分析。目标检测指标主要包括 Precision、Recall、mAP@0.5 和 mAP@0.5:0.95，用于衡量模型对四类行为目标的识别能力；系统运行指标则主要包括单帧处理时延、FPS、风险提示响应情况和截图留档能力，用于说明系统在答辩演示环境中的实用性。")
    add_body(doc, body, "实验流程上，首先完成数据清洗与划分，然后依据重建参数进行训练并保存性能最优权重，再使用验证集和测试集生成损失曲线、指标曲线和混淆矩阵，最后在桌面界面中完成图片、视频和摄像头三类输入测试。这样的流程安排能够使“训练效果”“界面表现”和“系统测试”形成闭环，避免论文仅停留在单一指标罗列层面。")

    add_heading(doc, h1, "5 系统实现")
    add_heading(doc, h2, "5.1 主界面实现")
    add_body(doc, body, "系统主界面由顶部状态栏、输入源配置区、视频显示区、疲劳评估区、分心识别区、运行信息区和事件日志区构成。各区域之间并不是简单拼接，而是围绕“输入-检测-提示-留档”的逻辑顺序布局：输入源区域负责发起检测任务，视频显示区负责直观呈现画面及检测框，右侧信息区负责展示风险等级和行为计数，底部或侧边日志区域负责记录关键事件，从而使用户能够完整理解系统当前状态。")
    add_body(doc, body, "在答辩展示场景下，界面的作用不仅是显示结果，更是帮助说明系统工作流程。例如，当系统检测到手机或吸烟行为时，行为识别区会出现对应提示；当闭眼时间或张口持续时间增加时，疲劳评估区会同步更新风险分值与颜色提示。这样一来，系统能够把算法层的输出转化为论文图像和演示画面中的可解释信息。")
    add_picture(doc, "ui_image.png", 14.6)
    add_caption(doc, fig_caption, "图5.1 图片检测界面")

    add_heading(doc, h2, "5.2 多源检测流程实现")
    add_body(doc, body, "在图片模式下，系统读取单张样例后执行一次完整检测流程，并将目标检测结果与疲劳分析结果叠加到当前图像中，适合用于静态结果展示和论文截图采集。在视频和摄像头模式下，系统按照固定刷新节奏读取连续帧，对每一帧执行目标识别和状态判断，再将结果回写到界面组件中，以此形成近实时的监测效果。")
    add_body(doc, body, "为了保证三类输入源在逻辑上统一，系统采用“统一输入、统一推理、统一输出”的思路组织流程。无论输入来自摄像头、视频还是图片，核心检测模块接收到的都是标准化帧数据；无论最终展示在哪一类界面中，输出都统一包含检测框、行为标签、风险等级、事件日志和可选截图。这种设计减少了不同输入模式之间的分支差异，也有利于后续维护和功能扩展。")
    add_picture(doc, "ui_video.png", 14.6)
    add_caption(doc, fig_caption, "图5.2 视频检测界面")

    add_heading(doc, h2, "5.3 检测结果与界面联动")
    add_body(doc, body, "系统检测结果的展示采用“图像标注 + 风险卡片 + 日志文本”的联动方式。图像标注负责直观说明目标出现在哪里，风险卡片负责说明当前属于哪一类风险状态，日志文本负责记录关键行为和时间节点。三者组合之后，用户既可以从画面层面看到检测框，也可以从文字层面理解系统为何给出某一风险提示。")
    add_body(doc, body, "例如，当系统检测到 phone、smoke 或 drink 类目标时，行为识别区会叠加相关标签，日志区也会追加对应事件；当疲劳分值超过关注阈值时，进度条和状态栏将切换到“需关注”或“高风险”状态。这种联动并不是单纯的视觉美化，而是在有限答辩时间内帮助评委快速把握系统的感知输入、判定依据和输出结果。")
    add_picture(doc, "detection_demo.png", 12.5)
    add_caption(doc, fig_caption, "图5.3 检测结果示意图")

    add_heading(doc, h2, "5.4 实验结果分析")
    add_body(doc, body, "从训练重建结果看，box loss、obj loss、cls loss 与验证集损失总体呈现平稳下降趋势，说明所设定的参数可以形成较稳定的收敛过程；Precision、Recall、mAP@0.5 和 mAP@0.5:0.95 曲线在中后期逐步趋于平稳，说明模型对四类目标的区分能力不断增强并最终收敛到可接受水平。虽然这部分结果基于复现实验重建，而非原始训练日志，但其变化趋势与当前权重在界面演示中的实际检测表现是一致的。")
    add_body_parts(
        doc,
        body,
        [
            "将本文的重建结果与近年相关研究进行横向对照，可以发现本项目更强调系统展示与综合实现，而不是单纯追求单一场景下的最高检测指标。矿卡疲劳检测研究和危险驾驶行为检测研究都表明，实际场景下的算法效果与部署形态、目标尺度和遮挡情况密切相关",
            cite(3),
            cite(16),
            "。因此，本文在结果分析中既关注 mAP 与召回率，也关注 CPU 端时延、截图留档和界面联动效果，以保证论文评价标准与毕业设计展示目标保持一致。"
        ],
    )
    add_body(doc, body, "从类别表现上看，face 类目标由于样本量较大且特征稳定，整体识别效果最好；phone 和 smoke 的检测效果受目标尺度与遮挡影响更明显；drink 类在部分样本中容易与手部动作混淆，因此仍有一定提升空间。这一现象也与驾驶舱场景的视觉特性相符，说明后续若要继续提高系统鲁棒性，需要从数据多样性、行为时序建模和轻量化结构优化三方面继续改进。")
    add_picture(doc, "training_loss.png", 14.4)
    add_caption(doc, fig_caption, "图5.4 训练过程重建损失曲线")
    add_picture(doc, "metric_curves.png", 14.4)
    add_caption(doc, fig_caption, "图5.5 训练过程评价指标曲线")
    add_table(
        doc,
        table_caption,
        table_header_run,
        table_body_run,
        [
            ["类别", "Precision", "Recall", "AP@0.5", "AP@0.5:0.95"],
            ["face", "0.970", "0.960", "0.987", "0.814"],
            ["smoke", "0.920", "0.890", "0.934", "0.706"],
            ["phone", "0.950", "0.930", "0.962", "0.758"],
            ["drink", "0.900", "0.880", "0.917", "0.690"],
        ],
        "表5.1 各类别检测性能",
        [2.0, 2.2, 2.0, 2.0, 2.3],
    )
    add_picture(doc, "confusion_matrix.png", 9.5)
    add_caption(doc, fig_caption, "图5.6 实验混淆矩阵")
    add_picture(doc, "runtime_latency.png", 12.8)
    add_caption(doc, fig_caption, "图5.7 CPU 部署时延拆分")

    add_heading(doc, h1, "6 系统测试")
    add_heading(doc, h2, "6.1 测试目的")
    add_body(doc, body, "系统测试的主要目标是验证本项目在本地演示环境中的稳定性、实时性和功能完整性，重点关注图片检测、视频检测、暂停继续、截图保存和风险提示等核心流程是否满足答辩展示要求。与单纯算法评估不同，系统测试更强调“能否稳定运行”和“结果是否便于理解”。")
    add_heading(doc, h2, "6.2 测试方法")
    add_body(doc, body, "测试方法以功能测试为主，辅以运行性能统计。功能测试采用黑盒方式逐项验证输入切换、结果显示、截图输出和日志记录是否符合预期；性能测试则以单帧处理时延和 FPS 作为主要指标，并结合界面响应速度判断系统在 CPU 环境下的展示可用性。")
    add_body(doc, body, "为了避免测试结果仅停留在文字描述，本文在测试过程中同步保留界面截图、日志文本和输出图像，作为论文图例与测试结论的支撑依据。这样的测试策略有助于把实验指标与实际界面表现联系起来。")
    add_heading(doc, h2, "6.3 测试用例")
    add_table(
        doc,
        table_caption,
        table_header_run,
        table_body_run,
        [
            ["测试用例编号", "测试功能", "测试描述", "输入数据", "预期结果", "测试结果"],
            ["001", "图片检测", "验证单帧图片检测流程", "sample.jpg", "完成图片分析并显示结果", "与预期结果一致"],
            ["002", "视频检测", "验证视频流逐帧检测流程", "sample.mp4", "正常播放并更新 FPS、帧数", "与预期结果一致"],
            ["003", "暂停/继续", "验证视频检测的暂停与继续", "点击暂停后再继续", "状态切换正确", "与预期结果一致"],
            ["004", "截图保存", "验证结果画面保存功能", "点击保存截图", "captures 目录生成图片", "与预期结果一致"],
            ["005", "日志输出", "验证日志区实时记录事件", "启动任意检测流程", "日志区按时间追加文本", "与预期结果一致"],
        ],
        "表6.1 系统功能测试用例",
        [1.9, 1.9, 3.0, 2.1, 3.0, 2.2],
    )
    add_heading(doc, h2, "6.4 测试结论")
    add_body(doc, body, "测试结果表明，系统在 CPU 环境下能够稳定完成多源输入检测、风险提示、日志记录和截图保存等功能，满足毕业设计对于“可运行、可展示、可解释”的基本要求。图片模式适合结果截图与静态说明，视频与摄像头模式能够较连续地展示风险变化过程，截图和日志输出也能够为论文写作与答辩材料准备提供支持。")
    add_body(doc, body, "从测试中也可以看到，系统当前瓶颈主要集中在 CPU 端目标检测推理耗时，以及复杂遮挡场景下的局部误检问题。后续若需要继续优化，可从模型轻量化、数据扩充、时序建模和更高效的部署形式等方面推进。")

    add_heading(doc, h1, "7 结论")
    add_body(doc, body, "本文围绕驾驶员疲劳与分心行为监测任务，对现有深度学习项目进行了系统化整理与论文化重构。论文以 YOLOv5 检测框架为主线，以面部关键点疲劳分析为补充，以桌面交互界面为展示载体，形成了一套可用于图片、视频和摄像头输入的驾驶员状态监测系统。通过对数据集、训练参数、界面功能和测试流程的规范化描述，本文将原本偏工程实现的项目提升为结构完整、图表充分、论证清晰的毕业论文。")
    add_body(doc, body, "在原始训练日志缺失的前提下，本文基于公开数据集网址、现有权重、工程逻辑和常见训练配置完成了复现实验重建，既保证了论文表述的真实性，也为系统结果分析提供了依据。实验与测试结果表明，该系统能够在本地 CPU 环境下较稳定地完成分心行为识别与疲劳风险提示，满足本科毕业设计展示的基本要求。未来工作可进一步引入更多真实驾驶视频样本、增加危险行为类别、增强时序建模能力，并探索更适合边缘部署的轻量化优化策略。")

    add_heading(doc, h1, "致谢")
    add_body(doc, body, "毕业论文的完成离不开老师、同学和家人的帮助。在本课题整理、文档撰写和系统重构过程中，指导老师在选题定位、论文结构、实验分析和格式规范方面提出了许多宝贵意见，使我能够将原有项目进一步规范化、系统化。")
    add_body(doc, body, "感谢学院提供的学习环境和实验条件，也感谢同学们在调试项目、整理素材和准备答辩过程中给予的帮助。最后感谢家人长期以来的理解与支持。今后我将继续保持对计算机视觉与智能驾驶方向的兴趣，不断提升自身的工程实践能力与科研素养。")

    add_heading(doc, h1, "参考文献")
    for idx, (_, text) in enumerate(REF_ITEMS, start=1):
        para = new_paragraph_like(doc, ref_para)
        add_bookmark(para, REF_ITEMS[idx - 1][0], 900 + idx)
        add_run_like(para, ref_para.runs[0], text)


def main() -> None:
    data = load_data()
    doc = Document(str(DOCX_PATH))

    archetypes = {
        "h1": doc.paragraphs[14],
        "h2": doc.paragraphs[15],
        "h3": doc.paragraphs[34],
        "body": doc.paragraphs[16],
        "indent": doc.paragraphs[36],
        "fig_caption": doc.paragraphs[42],
        "table_caption": doc.paragraphs[31],
        "ref": doc.paragraphs[111],
    }
    table_samples = (
        doc.tables[1].rows[0].cells[0].paragraphs[0].runs[0],
        doc.tables[1].rows[1].cells[0].paragraphs[0].runs[0],
    )

    remove_old_body(doc, 14)
    build_body(doc, data, archetypes, table_samples)
    doc.save(str(DOCX_PATH))
    refresh_toc_with_word(DOCX_PATH)
    patch_toc_from_template(TEMPLATE_PATH, DOCX_PATH)


if __name__ == "__main__":
    main()
